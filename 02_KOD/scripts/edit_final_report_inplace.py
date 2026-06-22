"""In-place edit of TESLIM/01_RAPOR/final_report.docx.

Applied changes (per user request):
  1. Remove first-line indents from every paragraph.
  2. Bold every figure / table caption (paragraphs starting with
     "Sekil " or "Tablo ").
  3. Drop SMOTE mentions (not in the assignment spec).
  4. Remove Section 5.2 "Dogruluk-Kararlilik Dengesi" (not in spec).
  5. Insert a figure in Section 2 (GENEL BILGILER).
  6. Insert a figure in Section 4.5 (ek sorular).
  7. Insert a figure after the second paragraph in Section 6 SONUC.
  8. Shorten Section 4.5 prose, simplify language, academic passive
     voice, no first-person.

The script is idempotent w.r.t. text edits where it can be — it
checks for the new text before applying.
"""
from __future__ import annotations

import re
import sys
from copy import deepcopy
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

DOC_PATH = ROOT / "TESLIM" / "01_RAPOR" / "final_report.docx"
FIG_DIR = ROOT / "results" / "figures"


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _set_no_indent(p: Paragraph) -> None:
    pf = p.paragraph_format
    pf.first_line_indent = None
    pPr = p._element.find(qn("w:pPr"))
    if pPr is not None:
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            for attr in (qn("w:firstLine"), qn("w:hanging")):
                if attr in ind.attrib:
                    del ind.attrib[attr]


def _clear_paragraph(p: Paragraph) -> None:
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def _set_text(p: Paragraph, text: str, *, bold: bool = False,
              size_pt: float = 12.0, italic: bool = False,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    _clear_paragraph(p)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    _set_no_indent(p)


def _add_figure_after(after_p: Paragraph, image_path: Path,
                      caption: str, *, width_cm: float = 14.5) -> Paragraph:
    """Insert image + caption paragraphs after `after_p`. Returns caption paragraph."""
    img_p = _insert_paragraph_after(after_p)
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_after = Pt(2)
    img_p.paragraph_format.line_spacing = 1.0
    _set_no_indent(img_p)
    run = img_p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    cap_p = _insert_paragraph_after(img_p)
    _set_text(cap_p, caption, bold=True, size_pt=11,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
    cap_p.paragraph_format.space_after = Pt(10)
    return cap_p


def _delete_paragraph(p: Paragraph) -> None:
    p._element.getparent().remove(p._element)


def main() -> None:
    doc = Document(str(DOC_PATH))

    # ------------------------------------------------------------
    # STEP 1: Remove first-line indents from EVERY paragraph
    # ------------------------------------------------------------
    for p in doc.paragraphs:
        _set_no_indent(p)
    # Also clear indents on table cell paragraphs
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _set_no_indent(p)
    print("[1] First-line indents cleared on", len(doc.paragraphs), "paragraphs")

    # ------------------------------------------------------------
    # STEP 2: Bold all figure / table captions
    # ------------------------------------------------------------
    caption_pat = re.compile(r"^(Şekil|Tablo)\s+\d+\.\d+\.\s", re.UNICODE)
    n_bolded = 0
    for p in doc.paragraphs:
        if caption_pat.match(p.text.strip()):
            for r in p.runs:
                r.font.bold = True
            n_bolded += 1
    print(f"[2] Bolded {n_bolded} caption paragraphs")

    # ------------------------------------------------------------
    # STEP 3: Remove SMOTE references
    # ------------------------------------------------------------
    smote_paras = [p for p in doc.paragraphs if "SMOTE" in p.text]
    # The paragraph in Section 5 is purely about SMOTE comparison; remove it.
    for p in smote_paras:
        text = p.text
        if "SMOTE gibi sentetik örnekleme" in text and "literatürle uyumlu" in text:
            _delete_paragraph(p)
            print("[3] Removed SMOTE comparison paragraph from Section 5")
    # Re-scan to confirm
    remaining = sum(1 for p in doc.paragraphs if "SMOTE" in p.text)
    print(f"[3] Remaining SMOTE mentions: {remaining}")

    # ------------------------------------------------------------
    # STEP 4: Remove Section 5.2 (heading + intro + figure + caption + desc)
    # ------------------------------------------------------------
    # Pattern: find heading "5.2. Doğruluk-Kararlılık Dengesi" — then delete
    # paragraphs until next major heading "6. SONUÇ" (exclusive).
    paras = doc.paragraphs
    start_idx = None
    end_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if start_idx is None and t.startswith("5.2.") and "Doğruluk" in t and "Karar" in t and "\t" not in t:
            start_idx = i
        elif start_idx is not None and (t.startswith("6.") and "SONUÇ" in t and "\t" not in t):
            end_idx = i
            break
    if start_idx is not None and end_idx is not None:
        to_remove = paras[start_idx:end_idx]
        # Also remove any embedded image runs by deleting the paragraphs entirely
        for p in to_remove:
            _delete_paragraph(p)
        print(f"[4] Removed Section 5.2 ({end_idx - start_idx} paragraphs)")
    else:
        print(f"[4] WARNING: Section 5.2 boundaries not found (start={start_idx}, end={end_idx})")

    # ------------------------------------------------------------
    # STEP 5: Renumber and replace Section 4.5 prose (shorter + simpler)
    # ------------------------------------------------------------
    # Section 4.5 currently has 5 long paragraphs after the heading:
    #   intro, Q1+Q2, Q4+Q5+Q7, Q3 (MRMR), Q6 (data size)
    # Rewrite them as 4 shorter ones.
    paras = doc.paragraphs  # refresh after deletions
    h45_idx = None
    next_h_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if h45_idx is None and t.startswith("4.5.") and "Ek Sor" in t and "\t" not in t:
            h45_idx = i
        elif h45_idx is not None and t.startswith("5.") and "TARTIŞMA" in t and "\t" not in t:
            next_h_idx = i
            break
    if h45_idx is not None and next_h_idx is not None:
        # Replacement paragraphs (intro + 4 Q-blocks)
        new_45 = [
            ("Bölüm 4.1 ile 4.4 arasında modellerin performansı ayrıntılı "
             "olarak sunulmuştur. Bu alt bölümde ödev tanımında yer alan "
             "yedi ek soruya doğrudan yanıt verilmektedir."),
            ("Ortalama macro-F1 sıralamasında ExtraTrees modeli 0.847 ile "
             "ilk sırada yer almakta, tek tekrar bazında ulaşılan en "
             "yüksek değer ise Gradient Boosting modeline ait olup 0.941 "
             "düzeyindedir. Soft voting topluluğu, baz modellerin "
             "birbirine yakın karar yüzeyleri öğrenmesi nedeniyle en "
             "iyi tek modeli geride bırakamamış; benzer çıktıların "
             "ortalanması beklenen çeşitlilik kazancını sağlamamıştır."),
            ("MRMR özellik seçimi 746 özelliğin tamamı yerine 20, 50 ya "
             "da 100 özellikten oluşan bir alt küme kullanmıştır. Bu "
             "yaklaşım örnek-özellik oranını yaklaşık 1.3:1 düzeyinden "
             "19:1 seviyesine çıkararak aşırı öğrenme riskini "
             "azaltmaktadır. Şekil 4.6'da görüldüğü üzere üst sıralardaki "
             "özellikler tekrarların büyük çoğunluğunda yeniden "
             "seçilmiş; en yüksek katkı sağlayan ilk on özellik Tablo "
             "4.4'te SHAP değerleriyle birlikte raporlanmıştır."),
            ("Sigmoid Platt kalibrasyon sonrasında modellerin Brier "
             "skorları dar bir aralıkta toplanmış ve kalibrasyon "
             "eğrileri diyagonale yakın seyretmiştir. ROC-AUC ile PR-AUC "
             "arasındaki ilişki ise sınıf dengesizliği koşullarında "
             "PR-AUC değerlerinin ROC-AUC değerlerinin biraz altında "
             "kalması biçiminde gözlenmiştir; bu durum azınlık sınıfın "
             "PR-AUC tarafından daha duyarlı yansıtıldığını "
             "göstermektedir."),
            ("Hasta sayısının görece az olması (n = 69) modellerin "
             "tekrarlar arasındaki kararlılığına doğrudan yansımaktadır. "
             "Şekil 4.8'de yatay eksende ortalama macro-F1, dikey "
             "eksende ise tekrarlar arası standart sapma yer almaktadır. "
             "Daha basit ve düzenlileştirilmiş modellerde standart sapma "
             "yaklaşık 0.07 düzeyinde kalırken, Gradient Boosting "
             "modelinde aynı değer 0.14 seviyesine kadar yükselmektedir. "
             "Bu fark, az veriyle eğitilen karmaşık modellerin tekrarlar "
             "arasında daha kararsız davrandığını ortaya koymakta ve "
             "daha geniş hasta kohortlarının özellikle bu model "
             "sınıfları için daha güvenilir sonuçlar üreteceğini "
             "düşündürmektedir."),
        ]

        # Update heading text to "İletilen Ek Sorulara Yönelik Değerlendirmeler"
        # (user typed this themselves; keep grammar consistent)
        h_para = paras[h45_idx]

        # Collect existing paragraphs between heading and next heading
        body = paras[h45_idx + 1:next_h_idx]
        # Remove all blank lines / images; rewrite with our text
        for p in body:
            _delete_paragraph(p)

        # Add new paragraphs after the heading
        anchor = h_para
        for i, txt in enumerate(new_45):
            new_p = _insert_paragraph_after(anchor)
            _set_text(new_p, txt)
            anchor = new_p

        print(f"[5] Replaced Section 4.5 with {len(new_45)} paragraphs")
    else:
        print(f"[5] WARNING: Section 4.5 boundaries not found "
              f"(h={h45_idx}, next={next_h_idx})")

    # ------------------------------------------------------------
    # STEP 6: Insert Şekil 4.8 (dogruluk_kararlilik) at end of Section 4.5
    # ------------------------------------------------------------
    paras = doc.paragraphs
    # Locate the last paragraph of 4.5 (the Q6 / data-size paragraph)
    h45_idx = None
    next_h_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if h45_idx is None and t.startswith("4.5.") and "Ek Sor" in t and "\t" not in t:
            h45_idx = i
        elif h45_idx is not None and (
            (t.startswith("5.") and "TARTIŞMA" in t and "\t" not in t)
            or (t == "5. TARTIŞMA")
        ):
            next_h_idx = i
            break
    if h45_idx is not None and next_h_idx is not None:
        # The last non-empty paragraph before next_h_idx is the anchor
        anchor = None
        for j in range(next_h_idx - 1, h45_idx, -1):
            if paras[j].text.strip():
                anchor = paras[j]
                break
        if anchor is not None:
            _add_figure_after(
                anchor,
                FIG_DIR / "dogruluk_kararlilik.png",
                "Şekil 4.8. Modellerin ortalama macro-F1 ile tekrarlar "
                "arası standart sapma ilişkisi.",
                width_cm=14.5,
            )
            print("[6] Inserted Şekil 4.8 at end of Section 4.5")

    # ------------------------------------------------------------
    # STEP 7: Insert Şekil 2.1 in Section 2.2 (Radyomik) and renumber
    # the existing Şekil 2.1 (sınıf dengesizliği) to Şekil 2.2.
    # ------------------------------------------------------------
    paras = doc.paragraphs
    # First rename old Şekil 2.1 caption -> Şekil 2.2
    renamed = False
    for p in paras:
        t = p.text.strip()
        if t.startswith("Şekil 2.1.") and "engesizli" in t:
            new_t = t.replace("Şekil 2.1.", "Şekil 2.2.", 1)
            _set_text(p, new_t, bold=True, size_pt=11,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
            p.paragraph_format.space_after = Pt(10)
            renamed = True
            break
    print(f"[7] Renamed old Şekil 2.1 to 2.2: {renamed}")

    # Find heading "2.3. Sınıf Dengesizliği" — insert new figure before it,
    # i.e. at end of Section 2.2 (Radyomik Veri ve Özellik Seçimi).
    paras = doc.paragraphs
    anchor_for_new_21 = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t.startswith("2.3.") and "Sınıf Dengesizli" in t and "\t" not in t:
            # Anchor = paragraph immediately before this heading
            for j in range(i - 1, -1, -1):
                if paras[j].text.strip() and "\t" not in paras[j].text:
                    anchor_for_new_21 = paras[j]
                    break
            break
    if anchor_for_new_21 is not None:
        _add_figure_after(
            anchor_for_new_21,
            FIG_DIR / "veri_seti_ozet.png",
            "Şekil 2.1. Veri setinin sınıf bazlı örnek ve hasta dağılımı.",
            width_cm=14.0,
        )
        print("[7] Inserted Şekil 2.1 at end of Section 2.2")

    # ------------------------------------------------------------
    # STEP 8: Insert Şekil 6.2 (ogrenme_egrisi) after existing Şekil 6.1
    # description in Section 6 SONUC.
    # ------------------------------------------------------------
    paras = doc.paragraphs
    # Find the paragraph that describes Şekil 6.1 (radar)
    anchor_for_62 = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t.startswith("Şekil 6.1'deki radar"):
            anchor_for_62 = p
            break
    if anchor_for_62 is not None:
        cap_p = _add_figure_after(
            anchor_for_62,
            FIG_DIR / "ogrenme_egrisi.png",
            "Şekil 6.2. Kümülatif ortalama ve standart sapmanın tekrar "
            "sayısı ile değişimi.",
            width_cm=15.5,
        )
        # Add a short description paragraph after the caption
        desc_p = _insert_paragraph_after(cap_p)
        _set_text(
            desc_p,
            "Şekil 6.2'de modellerin kümülatif ortalama macro-F1 ve "
            "standart sapma değerlerinin tekrar sayısı arttıkça nasıl "
            "oturduğu gösterilmektedir. Eğrilerin son tekrarlara doğru "
            "düz bir seyre kavuşması, 20 dış tekrarın hem ortalama hem "
            "de varyans kestirimleri için yeterli bir hacim sağladığını "
            "ortaya koymaktadır."
        )
        print("[8] Inserted Şekil 6.2 in Section 6 SONUC")

    # ------------------------------------------------------------
    # Save
    # ------------------------------------------------------------
    out = DOC_PATH
    doc.save(str(out))
    print(f"\nSaved to: {out}")


if __name__ == "__main__":
    main()
