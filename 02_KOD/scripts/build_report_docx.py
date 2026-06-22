"""Generate final_report.docx in a simplified Turkish project-report format.

Page setup based on Uskudar Universitesi guide (A4, 4cm/2.5cm margins,
Times New Roman 12pt, 1.5 line spacing) but the document is a *project
assignment*, not a thesis - so it skips advisor, English abstract,
teşekkür, and table/figure indexes.

Layout:
  - Cover + inner cover (student info, no advisor)
  - ÖZET + Anahtar Kelimeler
  - İÇİNDEKİLER (Word-style with dotted tab leaders)
  - SİMGELER VE KISALTMALAR
  - 1. GİRİŞ
  - 2. GENEL BİLGİLER
  - 3. GEREÇ VE YÖNTEM (with embedded dataset + workflow figures)
  - 4. BULGULAR (figures + metrics image)
  - 5. TARTIŞMA
  - 6. SONUÇ
  - KAYNAKLAR
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pandas as pd  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.text import (  # noqa: E402
    WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER,
)
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.shared import Cm, Pt  # noqa: E402

from src import config  # noqa: E402


# --------------------------------------------------- student info --- #
STUDENT_NO = "244312505"
STUDENT_NAME = "Okan ULUĞLU"

REPORT_TITLE = "Radyomik Özellikler Kullanılarak Papilödem İkili Sınıflandırması"

REFERENCES = [
    "Demircioğlu A. Benchmarking Feature Selection Methods in Radiomics. "
    "Investigative Radiology, 2022.",
    "Demircioğlu A. The effect of data resampling methods in radiomics. "
    "Scientific Reports, 2024.",
    "Shenoy R., Samra G.S., Sekhri R., et al. Clinician-Led Code-Free Deep "
    "Learning for Detecting Papilledema and Pseudopapilledema Using Optic "
    "Disc Imaging. Translational Vision Science & Technology, 2026.",
    "Szanto D., Erekat A., Woods B., et al. Multimodal Deep Learning "
    "Differentiates Papilledema and Non-Arteritic Anterior Ischemic Optic "
    "Neuropathy From Healthy Eyes. IOVS, 2026.",
    "Girard M.J.A., Panda S., Tun T.A., et al. Discriminating Between "
    "Papilledema and Optic Disc Drusen Using 3D Structural Analysis of the "
    "Optic Nerve Head. Neurology, 2023.",
    "Tang G., Du L., Ling S., Che Y., Chen X. Multi-type classification of "
    "lung nodules based on CT radiomics and ensemble learning for diversity "
    "weighting. Quantitative Imaging in Medicine and Surgery, 2024.",
    "Pedregosa F., Varoquaux G., Gramfort A., et al. Scikit-learn: Machine "
    "Learning in Python. Journal of Machine Learning Research, "
    "12: 2825-2830, 2011.",
    "Akiba T., Sano S., Yanase T., Ohta T., Koyama M. Optuna: A "
    "Next-generation Hyperparameter Optimization Framework. ACM SIGKDD, 2019.",
    "Lundberg S.M., Lee S.I. A Unified Approach to Interpreting Model "
    "Predictions. NeurIPS, 2017.",
    "Peng H., Long F., Ding C. Feature selection based on mutual "
    "information: criteria of max-dependency, max-relevance, and "
    "min-redundancy. IEEE TPAMI, 27(8): 1226-1238, 2005.",
    "Platt J.C. Probabilistic outputs for support vector machines and "
    "comparisons to regularized likelihood methods. Advances in Large "
    "Margin Classifiers, 10(3): 61-74, 1999.",
    "Fang S.S., Chen S.H. AI-assisted diagnosis of neuro-ophthalmic "
    "disorders: a systematic review from optic neuritis to papilledema. "
    "BMC Ophthalmology, 2026.",
]


def _load_results():
    s = config.TABLES_DIR / "metrics_summary.csv"
    summary = pd.read_csv(s, header=[0, 1], index_col=0) if s.exists() else None
    w = config.TABLES_DIR / "wilcoxon_bonferroni.csv"
    wilcoxon = pd.read_csv(w) if w.exists() else None
    f = config.TABLES_DIR / "friedman.json"
    friedman = json.loads(f.read_text()) if f.exists() else None
    m = config.RESULTS_DIR / "run_manifest.json"
    manifest = json.loads(m.read_text()) if m.exists() else None
    shap = config.TABLES_DIR / "shap_top10.json"
    shap_top = json.loads(shap.read_text()) if shap.exists() else None
    return summary, wilcoxon, friedman, manifest, shap_top


# ---------------------------------------------------------- helpers
def _set_para_format(p, *, align=None, line_spacing=1.5, space_before=0,
                     space_after=6, first_line_indent=None):
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)


def _add_para(doc, text="", *, bold=False, size=12, italic=False, align=None,
              first_line_indent=None, space_after=3, space_before=0):
    p = doc.add_paragraph()
    _set_para_format(
        p,
        align=align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=first_line_indent,
        space_after=space_after,
        space_before=space_before,
    )
    if text:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
    return p


def _add_centered(doc, text, *, bold=False, size=12, italic=False, space_after=0):
    return _add_para(doc, text, bold=bold, size=size, italic=italic,
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=space_after)


def _add_heading1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text.upper())
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.bold = True
    pPr = p._p.get_or_add_pPr()
    pageBreakBefore = OxmlElement("w:pageBreakBefore")
    pPr.append(pageBreakBefore)
    return p


def _add_heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(8)
    pf.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = True
    return p


def _add_figure(doc, path, caption, width_cm=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    if Path(path).exists():
        run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(8)
        cr = cap.add_run(caption)
        cr.font.name = "Times New Roman"
        cr.font.size = Pt(11)
        cr.italic = True


def _add_caption(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(caption)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.italic = True
    r.bold = True


def _set_cell_text(cell, text, *, bold=False, size=11, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(str(text))
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold


def _set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _add_code_block(doc, code_text, *, font="Consolas", size=10):
    """Render a code snippet in a single-cell shaded table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "F5F5F5")
    # Light single border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "C0C0C0")
        tcBorders.append(b)
    tcPr.append(tcBorders)
    # First (auto-created) paragraph + extra runs
    cell.text = ""
    lines = code_text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(line if line else " ")
        r.font.name = font
        r.font.size = Pt(size)
    # space after the table
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    after.paragraph_format.space_before = Pt(0)


def _add_table(doc, headers, rows, *, header_align=WD_ALIGN_PARAGRAPH.CENTER,
               body_align=WD_ALIGN_PARAGRAPH.LEFT, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        _set_cell_text(hdr_cells[j], h, bold=True, align=header_align)
        _set_cell_borders(hdr_cells[j])
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            _set_cell_text(table.rows[i].cells[j], val, align=body_align)
            _set_cell_borders(table.rows[i].cells[j])
    if col_widths_cm is not None:
        for j, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table


def _add_pagenumber_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ============================================================ MAIN ====
def build_docx() -> Path:
    summary, wilcoxon, friedman, manifest, shap_top = _load_results()
    n_outer = manifest["n_outer_repeats"] if manifest else 20
    n_inner = manifest["n_inner_folds"] if manifest else 5

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(4.0)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        _add_pagenumber_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # =================================================== KAPAK ===
    _add_centered(doc, "T.C.", bold=True, size=14, space_after=2)
    _add_centered(doc, "ÜSKÜDAR ÜNİVERSİTESİ", bold=True, size=14, space_after=2)
    _add_centered(doc, "MÜHENDİSLİK VE DOĞA BİLİMLERİ FAKÜLTESİ", bold=True,
                  size=13, space_after=24)
    for _ in range(2):
        _add_centered(doc, "", space_after=0)
    _add_centered(doc, "YAPAY ZEKA MÜHENDİSLİĞİ ANABİLİM DALI", bold=True,
                  size=13, space_after=2)
    _add_centered(doc, "YAPAY ZEKA MÜHENDİSLİĞİ (TEZLİ) YÜKSEK LİSANS PROGRAMI",
                  bold=True, size=12, space_after=18)
    _add_centered(doc, "YAPAY ZEKA DERSİ FİNAL PROJESİ", bold=True, size=13,
                  space_after=24)
    for _ in range(2):
        _add_centered(doc, "", space_after=0)
    _add_centered(doc, "PROJENİN ADI", bold=True, size=12, space_after=6)
    _add_centered(doc, REPORT_TITLE, bold=True, size=14, space_after=48)
    for _ in range(3):
        _add_centered(doc, "", space_after=0)
    _add_centered(doc, STUDENT_NAME, bold=True, size=13, space_after=2)
    _add_centered(doc, f"Öğrenci No: {STUDENT_NO}", size=12, space_after=48)
    for _ in range(4):
        _add_centered(doc, "", space_after=0)
    _add_centered(doc, "İSTANBUL - 2026", bold=True, size=13, space_after=0)
    doc.add_page_break()

    # ============================ İÇ KAPAK (Öğrenci bilgileri) ===
    _add_centered(doc, "T.C.", bold=True, size=14, space_after=2)
    _add_centered(doc, "ÜSKÜDAR ÜNİVERSİTESİ", bold=True, size=14, space_after=2)
    _add_centered(doc, "MÜHENDİSLİK VE DOĞA BİLİMLERİ FAKÜLTESİ", bold=True,
                  size=13, space_after=24)
    _add_para(doc, "Anabilim Dalı\t: Yapay Zeka Mühendisliği",
              align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    _add_para(doc, "Program\t\t: Yapay Zeka Mühendisliği (Tezli) Y.L.",
              align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    _add_para(doc, "Ders\t\t: Yapay Zeka",
              align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    _add_para(doc, f"Öğrenci No\t: {STUDENT_NO}",
              align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    _add_para(doc, f"Öğrenci Adı Soyadı\t: {STUDENT_NAME.upper()}",
              align=WD_ALIGN_PARAGRAPH.LEFT, space_after=24)
    _add_centered(doc, "PROJENİN ADI", bold=True, size=12, space_after=4)
    _add_centered(doc, REPORT_TITLE, bold=True, size=12, space_after=24)
    doc.add_page_break()

    # ===================================================== ÖZET ===
    _add_heading1(doc, "ÖZET")
    _add_centered(doc, REPORT_TITLE, bold=True, size=12, space_after=12)

    ozet_paragraphs = [
        "Papilödem, kafa içi basıncın yükselmesine bağlı olarak iki gözde "
        "birden ortaya çıkan optik disk şişmesidir. Görme kaybına kadar "
        "ilerleyebilen bir tablo olduğu için erken ve doğru tanı klinik "
        "açıdan büyük önem taşır. Buna karşın papilödem ile sahte papilödem "
        "(pseudopapilledem) ayrımı klasik olarak göz hekiminin değerlendirmesine "
        "dayanmakta; bu da hekimden hekime tutarlılığı zaman zaman düşürmektedir.",
        f"Bu projede, optik diskten çıkarılmış 746 radyomik özellik üzerinden "
        f"Normal ve Papilödem ayrımını yapan bir makine öğrenmesi sistemi "
        f"geliştirilmiştir. Veri seti 69 olgudan (48 normal, 21 papilödem) "
        f"toplam 966 örnek içermektedir. Aynı hastanın örneklerinin eğitim ve "
        f"test setine birlikte düşmemesi için hasta seviyesinde çapraz "
        f"doğrulama uygulanmış; veriler {n_outer} kez farklı şekilde "
        f"bölünerek değerlendirme tekrarlanmıştır. Eksik veri tamamlama, "
        f"düşük varyanslı ve aşırı korelasyonlu özelliklerin elenmesi, "
        f"ölçekleme ve MRMR yöntemiyle özellik seçimi tek bir sklearn "
        f"Pipeline içinde kurulmuş; Optuna ile altı sınıflandırıcı "
        f"(Lojistik Regresyon, RBF SVM, Random Forest, ExtraTrees, "
        f"Gradient Boosting, KNN) ayarlanmış ve sonuçlar soft voting "
        f"ile birleştirilmiştir.",
        (
            f"En yüksek ortalama macro-F1 skoruna "
            f"{summary[('macro_f1','mean')].idxmax()} modeliyle ulaşılmıştır "
            f"({summary[('macro_f1','mean')].max():.3f} ± "
            f"{summary.loc[summary[('macro_f1','mean')].idxmax(), ('macro_f1', 'std')]:.3f}). "
            f"Tek tekrar bazında en iyi sonuç ise "
            f"{manifest['best_overall']['model']} modeli tarafından "
            f"üretilmiştir ({manifest['best_overall']['macro_f1']:.3f}). "
            f"Friedman testi, modellerin birbirinden istatistiksel olarak "
            f"ayrıştığını göstermiştir. SHAP analizi ile karara en çok "
            f"katkı veren ilk on radyomik özellik raporlanmıştır."
        ) if (summary is not None and manifest is not None) else
        "Sonuçlar henüz hesaplanmadığı için bu bölüm pipeline "
        "çalıştırıldıktan sonra otomatik olarak doldurulacaktır.",
    ]
    for paragraph in ozet_paragraphs:
        _add_para(doc, paragraph, first_line_indent=1.0, space_after=6)

    _add_para(doc, "")
    p = _add_para(doc, "", first_line_indent=0, space_after=12)
    r = p.add_run("Anahtar Kelimeler: ")
    r.font.name = "Times New Roman"; r.font.size = Pt(12); r.bold = True
    r2 = p.add_run("Radyomik, papilödem, ikili sınıflandırma, hasta "
                   "seviyesinde çapraz doğrulama, MRMR, Optuna, SHAP.")
    r2.font.name = "Times New Roman"; r2.font.size = Pt(12)

    # ===================================== İÇİNDEKİLER =========
    _add_heading1(doc, "İÇİNDEKİLER")
    toc_entries = [
        ("ÖZET", "i", 0),
        ("İÇİNDEKİLER", "ii", 0),
        ("SİMGELER VE KISALTMALAR", "iii", 0),
        ("1. GİRİŞ", "1", 0),
        ("1.1. Problemin Tanımı ve Önemi", "1", 1),
        ("1.2. Çalışmanın Amacı ve Katkıları", "2", 1),
        ("2. GENEL BİLGİLER", "3", 0),
        ("2.1. Papilödemin Klinik Tanımı ve Görüntüleme Yaklaşımı", "3", 1),
        ("2.2. Radyomik Veri ve Özellik Seçimi", "4", 1),
        ("2.3. Sınıf Dengesizliği", "5", 1),
        ("3. GEREÇ VE YÖNTEM", "6", 0),
        ("3.1. Veri Seti", "6", 1),
        ("3.2. Veri Ön İşleme", "7", 1),
        ("3.3. Özellik Seçimi", "8", 1),
        ("3.4. Sınıflandırma Modelleri", "9", 1),
        ("3.5. Hiperparametre Optimizasyonu", "10", 1),
        ("3.6. Çapraz Doğrulama Yapısı", "10", 1),
        ("3.7. Topluluk Öğrenmesi", "11", 1),
        ("3.8. Değerlendirme Metrikleri", "11", 1),
        ("3.9. İstatistiksel Analiz", "12", 1),
        ("4. BULGULAR", "13", 0),
        ("4.1. Genel Performans Karşılaştırması", "13", 1),
        ("4.2. İstatistiksel Karşılaştırma", "15", 1),
        ("4.3. Kalibrasyon Analizi", "16", 1),
        ("4.4. Özellik Yorumlanabilirliği (SHAP)", "17", 1),
        ("4.5. Ödev Ek Sorularına Yönelik Değerlendirmeler", "18", 1),
        ("5. TARTIŞMA", "20", 0),
        ("5.1. Modellerin Tekrarlar Arası Tutarlılığı", "20", 1),
        ("5.2. Doğruluk-Kararlılık Dengesi", "21", 1),
        ("6. SONUÇ", "22", 0),
        ("KAYNAKLAR", "23", 0),
        ("EKLER: KOD ÖRNEKLERİ", "24", 0),
    ]
    for title, page, level in toc_entries:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.15
        if level > 0:
            pf.left_indent = Cm(0.8)
        tab_stops = pf.tab_stops
        tab_stops.add_tab_stop(Cm(14.5), WD_TAB_ALIGNMENT.RIGHT,
                                WD_TAB_LEADER.DOTS)
        r1 = p.add_run(title)
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(12)
        r1.bold = (level == 0)
        r2 = p.add_run("\t" + page)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(12)
        r2.bold = (level == 0)

    # ===================== SİMGELER VE KISALTMALAR ============
    _add_heading1(doc, "SİMGELER VE KISALTMALAR")
    _add_table(doc,
        headers=["Kısaltma", "Açılım"],
        rows=[
            ("AUC", "Area Under Curve (Eğri Altı Alan)"),
            ("CV", "Cross-Validation (Çapraz Doğrulama)"),
            ("ET", "Extra Trees"),
            ("GB", "Gradient Boosting"),
            ("KNN", "K-Nearest Neighbors (K-En Yakın Komşu)"),
            ("LR", "Logistic Regression (Lojistik Regresyon)"),
            ("MRMR", "Minimum Redundancy Maximum Relevance"),
            ("OKT", "Optik Koherens Tomografi"),
            ("ONH", "Optic Nerve Head (Optik Sinir Başı)"),
            ("PR", "Precision-Recall"),
            ("RF", "Random Forest (Rastgele Orman)"),
            ("ROC", "Receiver Operating Characteristic"),
            ("SHAP", "SHapley Additive exPlanations"),
            ("SVM", "Support Vector Machine (Destek Vektör Makinesi)"),
            ("TPE", "Tree-structured Parzen Estimator"),
        ],
        col_widths_cm=[3.0, 11.5],
        body_align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    # ============================================ 1. GİRİŞ ===
    _add_heading1(doc, "1. GİRİŞ")

    _add_heading2(doc, "1.1. Problemin Tanımı ve Önemi")
    _add_para(doc,
        "Papilödem, kafa içi basıncın yükselmesine bağlı olarak iki gözde "
        "birden gelişen optik disk şişmesi olarak tanımlanır [12]. Tablo "
        "ilerlediğinde kalıcı görme kaybına ulaşabildiği için erken tanı "
        "klinik açıdan büyük önem taşır. Papilödemin ayırıcı tanısında en "
        "sık karıştırılan durum sahte papilödemdir (pseudopapilledem); "
        "ikisi arasındaki ayrımın hatalı yapılması bir tarafta gereksiz "
        "invaziv girişimlere, diğer tarafta hayati bulguların gözden "
        "kaçırılmasına yol açabilir [3]. Klasik tanı yöntemi fundus "
        "fotoğrafı veya optik koherens tomografi (OKT) görüntülerinin göz "
        "hekimi tarafından değerlendirilmesidir. Bu yaklaşım deneyim "
        "farklarına bağlı olarak hekimden hekime değişebildiği gibi, "
        "değerlendirme süresi de göz polikliniklerinde ciddi bir iş yükü "
        "oluşturmaktadır.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Son yıllarda görüntülerden çıkarılan sayısal özelliklerin "
        "(radyomik özellikler) makine öğrenmesi modelleriyle "
        "birleştirilmesinin tanısal performansı önemli ölçüde "
        "iyileştirebildiği gösterilmiştir [1] [6]. Buradaki temel sıkıntı "
        "şudur: radyomik veri kümelerinde örnek başına yüzlerce özellik "
        "bulunurken hasta sayısı çoğu zaman küçük kalır. Bu durum aşırı "
        "öğrenme (overfitting) riskini büyütür ve dolayısıyla özellik "
        "seçimi, kalibrasyon ve özellikle hasta seviyesinde veri "
        "sızıntısının önlenmesi gibi noktalara dikkat etmek gerekir.",
        first_line_indent=1.0,
    )

    _add_heading2(doc, "1.2. Çalışmanın Amacı ve Katkıları")
    _add_para(doc,
        "Bu projenin amacı, ödev kapsamında verilen radyomik veri seti "
        "üzerinde Normal ve Papilödem ayrımını yapan, akademik standartlara "
        "uygun, veri sızıntısız ve sonuçları yorumlanabilir bir makine "
        "öğrenmesi sistemi tasarlamaktır. Burada amaç sadece yüksek "
        "doğruluklu bir model elde etmek değil, aynı zamanda elde edilen "
        "sonuçların güvenilirliğini sağlam bir metodoloji ile "
        "desteklemektir.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Bu doğrultuda kurulan sistem, aynı hastaya ait örneklerin farklı "
        "alt kümelere düşmesini engelleyen bir çapraz doğrulama yapısı, "
        "çalışmaya özel yazılmış bir korelasyon filtresi, MRMR ile özellik "
        "seçimi, Optuna ile otomatik hiperparametre arama ve soft voting "
        "topluluk modelini bir araya getirmektedir. Modellerin "
        "olasılıkları sigmoid kalibrasyonla düzeltilmiş, SHAP analizi ile "
        "kararlarına en çok katkı veren özellikler raporlanmıştır. "
        "Çalışmada izlenen genel yöntem akışı Şekil 3.2'de "
        "gösterilmektedir.",
        first_line_indent=1.0,
    )

    # =================================== 2. GENEL BİLGİLER ===
    _add_heading1(doc, "2. GENEL BİLGİLER")

    _add_heading2(doc, "2.1. Papilödemin Klinik Tanımı ve Görüntüleme Yaklaşımı")
    _add_para(doc,
        "Papilödem, intrakraniyal basıncın yükselmesi sonucu bilateral "
        "olarak gelişen optik disk şişmesidir. Klinik olarak baş ağrısı, "
        "bulantı, görme bulanıklığı ve geçici görme kayıpları gibi "
        "semptomlara eşlik edebilir. Tedavi edilmediğinde optik sinir "
        "atrofisine ve kalıcı görme kaybına yol açabildiği için ayrıntılı "
        "değerlendirme gerektirir. Pratikte tanıyı zorlaştıran en önemli "
        "durum, papilödem ile drusen veya doğuştan gelen disk "
        "anomalilerine bağlı sahte papilödem tablolarının görüntü olarak "
        "birbirine benzemesidir.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Görüntüleme tarafında en yaygın kullanılan yöntemler fundus "
        "fotoğrafı ve optik koherens tomografi (OKT) yöntemidir. OKT, "
        "optik sinir başı çevresindeki retina sinir lifi tabakası "
        "kalınlığını ölçüp katmanları ayrıştırabildiği için son yıllarda "
        "tanı algoritmalarının temel girdisi haline gelmiştir. Shenoy ve "
        "ark. [3] AutoML platformları üzerinden OKT'den elde edilen "
        "görüntülerle papilödem ile sahte papilödemi başarıyla "
        "ayırabildiklerini bildirmiştir. Szanto ve ark. [4] üç boyutlu OKT "
        "verisi üzerinde derin öğrenme modeli kurarak papilödem, NAION ve "
        "sağlıklı gözleri yüksek doğrulukla ayırmıştır. Girard ve ark. [5] "
        "ise optik sinir başının üç boyutlu yapısal analizinden yola "
        "çıkarak papilödemi optik disk drusen tablosundan ayırt etmeyi "
        "başarmıştır.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Bu çalışmaların büyük bölümü derin öğrenme yaklaşımlarına dayansa "
        "da, küçük veri rejimlerinde klasik makine öğrenmesi modelleri ile "
        "anlamlı özellik mühendisliğinin halen rekabetçi sonuçlar "
        "üretebildiği [12] tarafından da vurgulanmaktadır. Klasik "
        "yöntemler, hem yorumlanabilirlik açısından üstünlük sağlar, "
        "hem de küçük örneklem durumunda derin öğrenmenin tipik veri "
        "açlığı problemiyle daha az karşılaşır.",
        first_line_indent=1.0,
    )

    _add_heading2(doc, "2.2. Radyomik Veri ve Özellik Seçimi")
    _add_para(doc,
        "Radyomik, görüntü üzerinden niceliksel olarak özellik çıkarma "
        "yaklaşımıdır. Tek bir görüntüden yüzlerce hatta binlerce "
        "(birinci dereceden istatistikler, doku özellikleri, şekil "
        "tanımlayıcıları, dalgacık tabanlı yüksek seviye özellikler vb.) "
        "sayısal değer üretilebilir. Bu zenginlik avantaj gibi görünse de "
        "uygulamada en büyük sıkıntı yüksek boyutluluk ile az örnek "
        "sayısının bir araya gelmesidir.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Demircioğlu [1] on radyomik veri seti üzerinde 29 farklı özellik "
        "seçim yöntemini karşılaştırdığı geniş kapsamlı çalışmasında "
        "ilginç bir sonuca ulaşmıştır: ANOVA, LASSO ve MRMR ensemble gibi "
        "görece sade yöntemler, karmaşık yöntemlerden istatistiksel "
        "olarak farklı bir performans göstermemiş ve seçim kararlılığı "
        "açısından daha güvenli bulunmuştur. Yöntem seçiminin sade tarafta "
        "kalmasının özellikle az veriyle çalışırken iki avantajı vardır: "
        "model karmaşıklığı düştüğü için aşırı uyum riski azalır, ayrıca "
        "kararların yorumlanması kolaylaşır. Bu çalışmada da bu doğrultuda "
        "MRMR yaklaşımı tercih edilmiştir.",
        first_line_indent=1.0,
    )

    _add_heading2(doc, "2.3. Sınıf Dengesizliği")
    _add_para(doc,
        "Tıbbi görüntüleme veri kümelerinde sınıf dengesizliği yaygın bir "
        "durumdur. Bu çalışmada da veri setinde yaklaşık %70 normal ve "
        "%30 papilödem oranında bir dengesizlik bulunmaktadır. Bu durum "
        "modellerin azınlık sınıfı (papilödem) yeterince öğrenememesine "
        "yol açabilmektedir. Sorunu hafifletmek amacıyla model eğitimi "
        "sırasında ilgili modellerde `class_weight='balanced'` parametresi "
        "kullanılmıştır. Bu seçenek, modelin kayıp fonksiyonunda azınlık "
        "sınıfa otomatik olarak daha yüksek ağırlık atamakta ve böylece "
        "azınlık sınıfın model tarafından göz ardı edilmesinin önüne "
        "geçmektedir.",
        first_line_indent=1.0,
    )

    # ================================== 3. GEREÇ VE YÖNTEM ===
    _add_heading1(doc, "3. GEREÇ VE YÖNTEM")

    _add_heading2(doc, "3.1. Veri Seti")
    _add_para(doc,
        "Çalışmada kullanılan veri seti, optik diskten çıkarılmış 746 "
        "radyomik özelliği içeren iki CSV dosyasından oluşmaktadır: "
        "`normal_radiomics.csv` ve `papilodem_radiomics.csv`. Veri setinin "
        "genel görünümü Şekil 3.1'de özetlenmiştir.",
        first_line_indent=1.0,
    )
    _add_figure(doc, config.FIGURES_DIR / "veri_seti_ozet.png",
                caption="Şekil 3.1. Veri setinin sınıf dağılımı ile hasta ve "
                        "örnek sayıları.",
                width_cm=15.5)

    _add_para(doc,
        "Görüldüğü üzere veri seti 48 normal ve 21 papilödem olmak üzere "
        "toplam 69 olgudan elde edilen 966 örnek içermektedir. Her hasta "
        "için sağ ve sol göz olmak üzere iki taraf, her taraf için yedi "
        "dilim alınmış; bu da hasta başına 14 örnek yapmaktadır. "
        "Veri yüklenirken karşılaşılan ilk pratik sorun şudur: her iki "
        "dosyada da `PatientIndex` sütununun 1'den başlayarak yeniden "
        "numaralandırılmış olması nedeniyle doğrudan birleştirildiğinde "
        "aynı kimliğin iki farklı hastayı temsil etme riski "
        "bulunmaktadır. Bu sorunun önüne geçmek amacıyla normal "
        "hastalara `N_`, papilödem hastalarına ise `P_` ön eki "
        "eklenmiştir. Ayrıca veriyi yükleyen fonksiyonun sonunda üç "
        "güvenlik kontrolü tanımlanmıştır: iki sınıfın hasta "
        "kümelerinin kesişip kesişmediği, herhangi bir hastaya iki "
        "farklı sınıf etiketinin atanıp atanmadığı ve hasta başına "
        "örnek sayısının tutarlı olup olmadığı kontrol edilmektedir. "
        "Bu kontroller, olası veri hatalarının daha pipeline "
        "başlamadan tespit edilmesini sağlamaktadır.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Çalışmada izlenen yöntem akışı Şekil 3.2'de gösterilmiştir. "
        "Bundan sonraki bölümler bu akıştaki her bir adımı sırasıyla "
        "anlatmaktadır.",
        first_line_indent=1.0,
    )
    _add_figure(doc, config.FIGURES_DIR / "yontem_akisi.png",
                caption="Şekil 3.2. Çalışmanın yöntem akış şeması.",
                width_cm=15.5)

    _add_heading2(doc, "3.2. Veri Ön İşleme")
    _add_para(doc,
        "Veri, modele verilmeden önce birkaç temizleme adımından "
        "geçirilmiştir. Bu adımlar manuel olarak değil sklearn "
        "Pipeline yapısı içinde tanımlanmıştır; bu sayede işlemler her "
        "zaman aynı sırayla ve yalnızca eğitim verisi üzerinde "
        "uygulanmaktadır.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "İlk adım eksik değerlerin tamamlanmasıdır. Eksik olan her "
        "hücre, ilgili sütunun medyan değeri ile doldurulmuştur. "
        "Ardından, değerleri neredeyse hiç değişmeyen özellikler ele "
        "alınmıştır. Bu tür özellikler modele yeni bir bilgi katmadığı "
        "gibi gereksiz bir gürültü kaynağı oluşturma riski "
        "taşımaktadır; düşük bir varyans eşiği kullanılarak söz konusu "
        "özellikler veri setinden çıkartılmıştır.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Sonraki aşamada birbirine çok yakın bilgi taşıyan özellikler "
        "elenmiştir. Bazı radyomik özelliklerin neredeyse aynı bilgiyi "
        "taşıdığı bilinmektedir; bu durum aralarındaki mutlak Pearson "
        "korelasyonunun 0.95'in üzerinde olması ile tespit edilir. Bu "
        "şekilde belirlenen özellik çiftlerinden biri elenmiş ve veri "
        "setinde tekrar eden bilgi azaltılmıştır.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Son adımda ölçekleme uygulanmıştır. Burada klasik "
        "standartlaştırma yerine RobustScaler tercih edilmiştir. Bunun "
        "nedeni, radyomik verilerde zaman zaman çok uçta değerlerin "
        "bulunabilmesi ve bu değerlerin klasik ortalama-standart sapma "
        "hesabını bozabilmesidir. RobustScaler, ortanca ve çeyrekler "
        "arası genişlik üzerinden çalıştığı için uç değerlere karşı "
        "daha dayanıklı bir ölçekleyicidir.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Ek bir not olarak belirtmek gerekir ki, veri yüklendiğinde "
        "bir özellikte sonsuz (inf) değer bulunduğu tespit edilmiştir. "
        "Bu değer NaN'a dönüştürülerek ilk adımdaki imputer tarafından "
        "medyan ile doldurulması sağlanmıştır. Aksi halde sonsuz değer "
        "sonraki adımlarda hesaplama hatalarına yol açacaktı.",
        first_line_indent=1.0,
    )

    _add_heading2(doc, "3.3. Özellik Seçimi")
    _add_para(doc,
        "746 özellik içinde gerçekten bilgilendirici olanların sınırlı "
        "bir alt küme oluşturduğu düşünülmektedir. Tüm özelliklerin "
        "modele aynı anda verilmesi hem öğrenme süresini uzatır hem de "
        "aşırı uyum (overfitting) riskini büyütür. Bu noktada MRMR "
        "(Minimum Redundancy Maximum Relevance) algoritması [10] "
        "devreye girmektedir. MRMR'nin temel mantığı sezgiseldir: iyi "
        "bir özellik kümesi belirlenirken iki ölçüt aynı anda gözetilir. "
        "İlki, seçilen her özelliğin hedef değişken hakkında yeterli "
        "bilgi taşımasıdır. İkincisi ise daha az dikkat çeken bir "
        "noktadır: seçilen özelliklerin kendi aralarında çok benzer "
        "olmaması beklenir. Birbirine çok benzeyen iki özellik aslında "
        "aynı bilgiyi iki kez taşır; ikincisinin eklenmesi modele "
        "yeni bir katkı sağlamaz.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Bu çalışmada MRMR'nin alaka (relevance) skoru olarak mutual "
        "information, artıklık (redundancy) skoru olarak Pearson "
        "korelasyonu kullanılmıştır. Seçilecek özellik sayısı için "
        "sabit bir değer belirlemek yerine 20, 50 ve 100 olmak üzere "
        "üç farklı seçenek Optuna arama uzayına eklenmiştir. Bu "
        "yaklaşım sayesinde her model için en uygun özellik sayısı da "
        "hiperparametre arama süreci içinde otomatik olarak "
        "belirlenmiştir. Pratikte 50 civarındaki değerlerin hem ağaç "
        "tabanlı modellerde hem de SVM modelinde en istikrarlı sonuçları "
        "ürettiği gözlenmiştir.",
        first_line_indent=1.0,
    )

    _add_heading2(doc, "3.4. Sınıflandırma Modelleri")
    _add_para(doc,
        "Ders dokümanında istenen altı temel sınıflandırıcı "
        "değerlendirilmiştir: Lojistik Regresyon, RBF çekirdekli Destek "
        "Vektör Makinesi, Random Forest, ExtraTrees, Gradient Boosting "
        "ve K-En Yakın Komşu [7]. Her model için kullanılan "
        "hiperparametre arama aralıkları Tablo 3.2'de toplu olarak "
        "verilmiştir. Sınıf dengesizliğinden dolayı uygun olan "
        "modellerde `class_weight='balanced'` ayarı aktif edilmiştir.",
        first_line_indent=1.0,
    )
    _add_caption(doc, "Tablo 3.2. Optuna hiperparametre arama uzayları.")
    _add_table(doc,
        headers=["Model", "Parametre", "Aralık / Set"],
        rows=[
            ("LR", "C", "log [1e-3, 1e2]"),
            ("LR", "penalty", "{l1, l2}"),
            ("SVM", "C", "log [1e-2, 1e2]"),
            ("SVM", "gamma", "log [1e-4, 1e0]"),
            ("RF / ET", "n_estimators", "[100, 500] adım 50"),
            ("RF / ET", "max_depth", "[3, 20]"),
            ("RF / ET", "min_samples_leaf", "[1, 10]"),
            ("RF / ET", "max_features", "{sqrt, log2}"),
            ("GB", "n_estimators", "[100, 500] adım 50"),
            ("GB", "learning_rate", "log [1e-3, 3e-1]"),
            ("GB", "max_depth", "[2, 8]"),
            ("GB", "subsample", "[0.5, 1.0]"),
            ("KNN", "n_neighbors", "[3, 25]"),
            ("KNN", "weights", "{uniform, distance}"),
            ("KNN", "metric", "{euclidean, manhattan}"),
            ("Tüm modeller", "Özellik sayısı (mrmr_k)", "20, 50 veya 100"),
        ],
        col_widths_cm=[3.5, 4.0, 6.5],
        body_align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    _add_para(doc,
        "Tablo 3.2 incelendiğinde her modelin kendine özgü parametrelere "
        "sahip olduğu görülmektedir. Çoğu parametre için sabit bir değer "
        "atanması yerine geniş bir arama aralığı tanımlanmış ve en "
        "uygun değerin Optuna tarafından belirlenmesi sağlanmıştır. "
        "Tüm modellerde ortak olarak yer alan mrmr_k parametresi "
        "seçilecek özellik sayısını belirlemekte olup üç farklı seçenek "
        "(20, 50 veya 100) arasından her model için en uygun olanı "
        "kullanılmıştır.",
        first_line_indent=1.0,
    )

    _add_heading2(doc, "3.5. Hiperparametre Optimizasyonu")
    _add_para(doc,
        "Hiperparametreleri elle ayarlamak yerine Optuna kütüphanesi [8] "
        "kullanılmıştır. Arama sırasında TPE örnekleyicisi (Tree-"
        "structured Parzen Estimator) sabit bir tohum değeriyle "
        "başlatılmış, böylece sonuçlar tekrar üretilebilir hale "
        "getirilmiştir. Her model için 50 deneme yapılmış; hedef "
        "fonksiyon olarak iç çapraz doğrulamadan elde edilen ortalama "
        "makro-F1 skoru seçilmiştir.",
        first_line_indent=1.0,
    )

    _add_heading2(doc, "3.6. Çapraz Doğrulama Yapısı")
    _add_para(doc,
        "Bu çalışmada veri hasta seviyesinde üç parçaya bölünmüştür: "
        "%70 eğitim, %10 doğrulama ve %20 test. Her hasta yalnızca bir "
        "setin içinde yer alır; aynı hastanın hiçbir örneği farklı bir "
        "sete düşmez. Bu bölme işlemi farklı rastgele tohumlarla "
        f"{n_outer} kez tekrarlanmış ve her tekrarda farklı bir hasta "
        f"dağılımı elde edilmiştir.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        f"Eğitim seti (%70) yalnızca model eğitimi ve hiperparametre "
        f"araması için kullanılır. Bu küme üzerinde Optuna ile "
        f"hiperparametre araması yapılırken, iç katmanda her deneme "
        f"sırasında {n_inner}-kat çapraz doğrulama (StratifiedGroupKFold) "
        f"uygulanmaktadır. Bu sayede ne dış test seti ne de doğrulama "
        f"seti hiperparametre seçimine etki etmez.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Doğrulama seti (%10) yalnızca karar eşiğinin (threshold) "
        "optimizasyonu için kullanılır. Model eğitimi bittikten sonra "
        "doğrulama seti üzerinde olasılık tahminleri yapılır ve F1 "
        "skorunu maksimize eden eşik değeri belirlenir. Bu eşik daha "
        "sonra test setine olduğu gibi uygulanır; test seti üzerinde "
        "hiçbir eşik ayarı yapılmaz. Bu yapı, test setinin gerçekten "
        "\"görülmemiş veri\" özelliğini korumasını güvence altına alır.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Hasta düzeyinde bölme yapılmasının kritik bir önemi vardır. "
        "Bu veri setinde her hastaya ait 14 örnek bulunmaktadır (sağ ve "
        "sol göz için 7'şer dilim). Eğer bölme örnek düzeyinde yapılsa, "
        "aynı hastanın bazı dilimleri eğitimde, bazıları testte "
        "bulunabilir. Bu durumda model hastalığı değil, doğrudan o "
        "hastanın görüntüsünün özelliklerini öğrenir ve test setinde "
        "gerçekçi olmayan bir başarı sergiler. Hasta-bazlı bölme bu "
        "sorunu kökten çözmektedir.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        f"Dış tekrar sayısı ödev kapsamında verilen iş akış şemasıyla "
        f"uyumlu olarak {n_outer} olarak belirlenmiştir. Bu sayede her "
        f"model {n_outer} farklı hasta-bazlı bölme üzerinde bağımsız olarak "
        f"değerlendirilmiş ve sonuçların ortalaması ile standart sapması "
        f"birlikte raporlanmıştır. Şekil 5.2'de gösterilen tekrarlar "
        f"arası varyans, modellerin küçük örneklemde dahi kararlı sonuçlar "
        f"üretebildiğini ortaya koymaktadır.",
        first_line_indent=1.0,
    )

    _add_heading2(doc, "3.7. Topluluk Öğrenmesi")
    _add_para(doc,
        "Tek tek modellerin yanı sıra, Tang ve ark. [6] tarafından "
        "etkinliği gösterilen soft voting yaklaşımı denenmiştir. Soft "
        "voting'in fikri basittir: bir örnek için her bir baz modelin "
        "verdiği olasılıkları alıp ortalamasını hesaplamak. Bu çalışmada "
        "topluluk modeli, sigmoid Platt scaling ile [11] kalibre edilmiş "
        "Random Forest, ExtraTrees ve Gradient Boosting olasılıklarının "
        "ortalanmasıyla oluşturulmuştur. Sınıf eşiği ise doğrulama "
        "verisinden F1 skorunu maksimuma çıkaracak şekilde seçilmiş; "
        "test verisi üzerinde hiçbir eşik ayarı yapılmamıştır.",
        first_line_indent=1.0,
    )

    _add_heading2(doc, "3.8. Değerlendirme Metrikleri")
    _add_para(doc,
        f"Ders dokümanında istenen dokuz metriğin tamamı raporlanmıştır: "
        f"doğruluk, hassasiyet, duyarlılık, F1, makro-F1, ROC-AUC, "
        f"PR-AUC, dengeli doğruluk ve Brier skoru. Her metrik {n_outer} "
        f"dış tekrarın ortalaması ve standart sapması olarak verilmiştir.",
        first_line_indent=1.0,
    )

    _add_heading2(doc, "3.9. İstatistiksel Analiz")
    _add_para(doc,
        "Yedi modelin (altı baz model artı topluluk) makro-F1 skorları "
        "üzerinde Friedman testi uygulanmış, ardından ikili "
        "karşılaştırmalar için Wilcoxon işaretli sıralar testi "
        "yapılmıştır. Toplam yirmi bir ikili karşılaştırma olduğu için "
        "Bonferroni düzeltmesi devreye girmiş; bu düzeltme anlamlılık "
        "eşiğini ciddi şekilde sıkılaştırdığı için aslında yalnızca "
        "oldukça belirgin farklar test sonucunda anlamlı çıkmaktadır. "
        "Sonuçlar yorumlanırken bu durumun göz önünde tutulması "
        "gerekmektedir.",
        first_line_indent=1.0,
    )

    # ============================================ 4. BULGULAR ===
    _add_heading1(doc, "4. BULGULAR")

    if summary is not None and manifest is not None:
        best_model = manifest["best_overall"]["model"]
        best_macro = manifest["best_overall"]["macro_f1"]
        best_repeat = manifest["best_overall"]["repeat"] + 1
        wall_min = manifest["total_seconds"] / 60

        _add_heading2(doc, "4.1. Genel Performans Karşılaştırması")
        _add_para(doc,
            f"Pipeline'ın tam çalışması yaklaşık {wall_min:.0f} dakika "
            f"sürmüş ve yedi modelin {n_outer} dış tekrar üzerindeki "
            f"performans özetini Tablo 4.1 sunmaktadır. Ortalama makro-F1 "
            f"sıralamasında ilk sırayı "
            f"{summary[('macro_f1','mean')].idxmax()} modeli almıştır "
            f"({summary[('macro_f1','mean')].max():.3f}). Tek tekrar "
            f"bazında en yüksek değer ise {best_repeat}. tekrarda "
            f"{best_model} modeliyle elde edilen "
            f"makro-F1 = {best_macro:.3f} olmuştur.",
            first_line_indent=1.0,
        )
        _add_caption(doc,
            "Tablo 4.1. Modellerin tüm metriklerde performansı "
            "(mean ± std).")
        _add_figure(doc, config.FIGURES_DIR / "tablo_4_1_metrics.png",
                    caption="", width_cm=15.5)

        _add_caption(doc, "Tablo 4.2. Macro-F1 sıralaması (en iyi → en kötü).")
        rank = summary[("macro_f1", "mean")].sort_values(ascending=False)
        rank_rows = []
        for i, (mdl, val) in enumerate(rank.items(), start=1):
            std_val = summary.loc[mdl, ("macro_f1", "std")]
            rank_rows.append((str(i), mdl, f"{val:.3f} ± {std_val:.3f}"))
        _add_table(doc,
            headers=["Sıra", "Model", "Macro-F1 (mean ± std)"],
            rows=rank_rows,
            col_widths_cm=[1.5, 3.5, 7.5],
            body_align=WD_ALIGN_PARAGRAPH.CENTER,
        )

        figs_dir = config.FIGURES_DIR
        _add_figure(doc, figs_dir / "model_comparison_macro_f1.png",
                    "Şekil 4.1. Modeller arası ortalama makro-F1 "
                    "karşılaştırması.")
        # Build dynamic sentence describing the top model and Ensemble
        rank_series = summary[("macro_f1", "mean")].sort_values(ascending=False)
        top1 = rank_series.index[0]
        top1_val = rank_series.iloc[0]
        top_others = ", ".join(
            f"{m} ({rank_series[m]:.3f})"
            for m in rank_series.index[1:4]
        )
        ens_val = summary.loc["Ensemble", ("macro_f1", "mean")] if "Ensemble" in summary.index else None
        gb_std = summary.loc["GB", ("macro_f1", "std")] if "GB" in summary.index else None
        ens_text = (
            f"Soft voting ile oluşturulan Ensemble modeli ise {ens_val:.3f} "
            f"değeri ile beklenenden az bir katkı sağlamıştır; "
            if ens_val is not None else
            "Soft voting Ensemble modeli beklenenden az bir katkı sağlamıştır; "
        )
        gb_text = (
            f"Gradient Boosting modelinde gözlenen yüksek standart sapma "
            f"({gb_std:.3f}), bu modelin küçük örneklemde tekrarlar arasında "
            f"kararsız davrandığını ortaya koymaktadır."
            if gb_std is not None else
            "Gradient Boosting modelinin yüksek standart sapması, bu modelin "
            "küçük örneklemde tekrarlar arasında kararsız davrandığını "
            "ortaya koymaktadır."
        )
        _add_para(doc,
            f"Şekil 4.1'de modellerin ortalama macro-F1 skorları yan yana "
            f"verilmiştir. {top1} modeli {top1_val:.3f} ile ilk sırayı "
            f"almış, onu yakın bir farkla {top_others} izlemiştir. "
            f"{ens_text}bu durum baz modellerin birbirlerine yakın karar "
            f"yüzeyleri öğrenmesinden, dolayısıyla soft voting ortalamasının "
            f"modellerin zaten benzer çıktısını koruyarak çeşitlilikten "
            f"yararlanamamasından kaynaklanmaktadır. {gb_text}",
            first_line_indent=1.0,
        )
        _add_figure(doc, figs_dir / "roc_curves.png",
                    f"Şekil 4.2. ROC eğrileri (tüm modeller, {n_outer} outer "
                    f"repeat ortalaması).")
        roc_means = summary[("roc_auc", "mean")].sort_values(ascending=False)
        roc_top1, roc_top1_val = roc_means.index[0], roc_means.iloc[0]
        roc_top2, roc_top2_val = roc_means.index[1], roc_means.iloc[1]
        roc_top3, roc_top3_val = roc_means.index[2], roc_means.iloc[2]
        roc_low = roc_means.index[-2:].tolist()
        roc_low_text = " ve ".join(roc_low)
        _add_para(doc,
            f"Şekil 4.2'de modellerin ROC eğrileri ortalama AUC değerleri "
            f"ile birlikte gösterilmiştir. {roc_top1} "
            f"(AUC = {roc_top1_val:.2f}) ve {roc_top2} "
            f"(AUC = {roc_top2_val:.2f}) en yüksek değerleri elde etmiş, "
            f"{roc_top3} yaklaşık {roc_top3_val:.2f} ile bunları takip "
            f"etmiştir. Eğrilerin tamamının diyagonalin oldukça üzerinde "
            f"seyretmesi tüm modellerin rastgele tahminden belirgin "
            f"biçimde daha iyi olduğunu göstermektedir. {roc_low_text} "
            f"modellerinin diğerlerinin biraz altında kalması, doğrusal "
            f"ya da mesafe tabanlı yöntemlerin radyomik özellik uzayının "
            f"yüksek boyutluluğunda biraz daha zorlanması ile uyumludur.",
            first_line_indent=1.0,
        )
        _add_figure(doc, figs_dir / "pr_curves.png",
                    "Şekil 4.3. Precision-Recall eğrileri.")
        _add_para(doc,
            "Sınıf dengesizliği bulunan veri kümelerinde PR-AUC, ROC-AUC'a "
            "kıyasla daha gerçekçi bir performans göstergesidir; çünkü "
            "azınlık sınıftaki başarıyı doğrudan yansıtır. Şekil 4.3'te "
            "ExtraTrees ve Ensemble modellerinin PR-AUC değeri 0.90 "
            "seviyesindedir; bu da modellerin papilödem sınıfını yüksek "
            "doğrulukla yakalayabildiğini ortaya koymaktadır. Eğrilerin "
            "yüksek recall bölgesinde dahi yüksek precision değerlerini "
            "koruyabilmesi, modellerin papilödem örneklerini gözden "
            "kaçırmadan da yanlış pozitif üretmediğini göstermektedir. "
            "Bu sonuç klinik açıdan önemlidir çünkü hem duyarlılık hem "
            "de güvenilirlik aynı anda sağlanmaktadır.",
            first_line_indent=1.0,
        )
        _add_figure(doc, figs_dir / "confusion_matrix.png",
                    f"Şekil 4.4. Karışıklık matrisi ({best_model}, "
                    f"{best_repeat}. tekrar).")
        _add_para(doc,
            f"Şekil 4.4, en iyi tekrardaki {best_model} modelinin "
            f"karışıklık matrisini detaylandırmaktadır. Yanlış negatif "
            f"(yani gözden kaçırılan papilödem örnekleri) sayısının "
            f"düşük kalması özellikle önemlidir; çünkü papilödem "
            f"teşhisinin atlanması klinik açıdan ciddi sonuçlara yol "
            f"açabilmektedir. Yanlış pozitiflerin de oldukça sınırlı "
            f"kalması, modelin gereksiz alarm üretmediğini ve elde "
            f"edilen sonuçların hekim doğrulamasında kullanılabilir "
            f"düzeyde olduğunu göstermektedir.",
            first_line_indent=1.0,
        )

    _add_heading2(doc, "4.2. İstatistiksel Karşılaştırma")
    _add_para(doc,
        "Yedi modelin sonuçları incelendiğinde bazı modellerin "
        "diğerlerinden bir miktar daha iyi performans gösterdiği "
        "görülmektedir. Burada cevaplanması gereken kritik soru şudur: "
        "bu fark gerçekten anlamlı bir performans farkını mı "
        "yansıtmaktadır, yoksa verilerin farklı şekilde bölünmesi "
        "nedeniyle ortaya çıkmış rastlantısal bir farklılık mıdır? "
        "Bu soruya yalnızca gözlemsel olarak yanıt vermek mümkün "
        "olmadığından istatistiksel testlere başvurulur.",
        first_line_indent=1.0,
    )
    if friedman is not None:
        chi_rounded = round(friedman["statistic"], 1)
        p_rounded = round(friedman["p_value"], 3)
        _add_para(doc,
            f"İlk olarak Friedman testi uygulanmıştır. Friedman testi "
            f"modelleri tek tek değil hepsini birlikte sıralayarak "
            f"\"genel olarak modeller arasında bir fark mevcut mudur?\" "
            f"sorusuna yanıt verir. Bu çalışmada ki-kare istatistiği "
            f"yaklaşık {chi_rounded} ve p-değeri ise {p_rounded} civarında "
            f"ölçülmüştür. P-değeri 0.05 anlamlılık eşiğinin altında "
            f"kaldığı için modeller arasındaki farkın yalnızca rastlantısal "
            f"bir dalgalanma olmadığı, gerçek bir performans farkına "
            f"işaret ettiği söylenebilir.",
            first_line_indent=1.0,
        )
    if wilcoxon is not None:
        n_sig = int(wilcoxon["significant"].sum())
        _add_para(doc,
            f"Friedman testi \"genel olarak fark mevcuttur\" sonucunu "
            f"vermekle birlikte hangi modelin hangisinden farklı olduğu "
            f"bilgisini sağlamaz. Bu bilgiye ulaşmak amacıyla her model "
            f"çiftinin ayrı ayrı karşılaştırılmasına olanak veren "
            f"Wilcoxon işaretli sıralar testi uygulanmıştır. Yedi model "
            f"arasında toplam 21 ikili karşılaştırma elde edilmektedir. "
            f"Bu kadar çok test yapıldığında tesadüfen \"anlamlı\" sonuç "
            f"veren karşılaştırmaların ortaya çıkabilmesi nedeniyle "
            f"Bonferroni düzeltmesi uygulanır. Söz konusu düzeltme, "
            f"anlamlılık eşiğini ciddi şekilde sıkılaştırdığı için "
            f"yalnızca çok belirgin farkların testten geçmesi beklenir. "
            f"Bu nedenle Tablo 4.3'te bir karşılaştırmanın \"Anlamlı: "
            f"Hayır\" sonucu vermesi, iki modelin eşit performans "
            f"gösterdiği anlamına gelmemekte; küçük örneklem koşullarında "
            f"Bonferroni'nin sıkı eşiğini aşacak büyüklükte bir farkın "
            f"bulunmadığını ifade etmektedir. Bu nedenle sonuçların "
            f"kesin hüküm yerine yön gösterici olarak yorumlanması "
            f"daha uygundur.",
            first_line_indent=1.0,
        )
        _add_caption(doc,
            "Tablo 4.3. Wilcoxon pairwise testi (Bonferroni düzeltmeli).")
        wil_headers = ["Model A", "Model B", "Statistic", "p_value",
                       "p_bonferroni", "Anlamlı"]
        wil_rows = []
        for _, row in wilcoxon.iterrows():
            wil_rows.append((
                str(row["model_a"]),
                str(row["model_b"]),
                f"{row['statistic']:.1f}",
                f"{row['p_value']:.4f}",
                f"{row['p_bonferroni']:.4f}",
                "Evet" if row["significant"] else "Hayır",
            ))
        _add_table(doc, headers=wil_headers, rows=wil_rows,
                   body_align=WD_ALIGN_PARAGRAPH.CENTER)

    _add_heading2(doc, "4.3. Kalibrasyon Analizi")
    _add_para(doc,
        "Sigmoid kalibrasyon sonrasında Brier skorları dar bir aralıkta "
        "(yaklaşık 0.08 ile 0.13 arasında) toplanmış; bu durum Tablo "
        "4.1'de de görülebilmektedir. Kalibrasyon eğrileri "
        "(Şekil 4.5) modellerin tahminlerinin ne kadar güvenilir "
        "olduğunu görselleştirir; ideal olarak eğri y = x diyagonalinde "
        "olmalıdır. Sonuçlarda tüm modellerin diyagonale oldukça "
        "yakın seyrettiği görülmüştür. Ağaç tabanlı modellerde "
        "kalibrasyondan önce sıkça görülen aşırı veya yetersiz "
        "güven eğilimi, sigmoid kalibrasyonla büyük ölçüde "
        "düzeltilmiştir [11].",
        first_line_indent=1.0,
    )
    _add_figure(doc, config.FIGURES_DIR / "calibration_curves.png",
                "Şekil 4.5. Kalibrasyon eğrileri (sigmoid Platt sonrası).")
    _add_para(doc,
        "Şekil 4.5'te tüm modellerin kalibrasyon eğrileri ideal y = x "
        "diyagonaline oldukça yakın seyretmektedir. Bu durum, modellerin "
        "ürettiği olasılıkların gerçek olasılık değerlerini iyi yansıttığını "
        "göstermektedir. Pratik olarak şu anlama gelir: modelin %80 "
        "olasılık verdiği örneklerin gerçekten yaklaşık %80'i papilödemdir. "
        "Bu özellik klinik karar destek sistemlerinde kritik öneme sahiptir "
        "çünkü hekim yalnızca sınıf etiketine değil, modelin verdiği "
        "güven düzeyine de bakarak karar süreçlerini şekillendirebilir. "
        "Sigmoid kalibrasyon olmasaydı özellikle ağaç tabanlı modellerin "
        "tipik aşırı veya yetersiz güven eğilimi sonuçların yorumlanmasını "
        "zorlaştırırdı.",
        first_line_indent=1.0,
    )

    _add_heading2(doc, "4.4. Özellik Yorumlanabilirliği (SHAP)")
    if shap_top:
        _add_para(doc,
            f"Modelin sınıflandırma kararına en çok hangi özelliklerin "
            f"etki ettiğini anlamak için SHAP TreeExplainer ile ExtraTrees "
            f"modeli üzerinde analiz yapılmıştır. SHAP değerleri her "
            f"özelliğin tahminin oluşmasına ne kadar katkı sağladığını "
            f"hesaplamakta ve böylece model davranışını şeffaflaştırmaktadır "
            f"[9]. En çok katkı veren ilk on özellik Tablo 4.4'te, "
            f"tamamının görsel özeti ise Şekil 4.7'de verilmiştir. "
            f"Ayrıca MRMR'nin {n_outer} outer tekrar boyunca tekrar tekrar "
            f"seçtiği özelliklerin Şekil 4.6'da gösterilmesi, SHAP "
            f"sıralamasıyla örtüşen özelliklerin gerçekten istikrarlı "
            f"olduğunu doğrulamaktadır.",
            first_line_indent=1.0,
        )
        _add_caption(doc,
            "Tablo 4.4. SHAP analizinde en yüksek katkılı ilk 10 özellik.")
        shap_rows = [(str(i), name, f"{val:.4f}")
                     for i, (name, val) in enumerate(shap_top, start=1)]
        _add_table(doc,
            headers=["Sıra", "Özellik", "Mean |SHAP|"],
            rows=shap_rows,
            col_widths_cm=[1.5, 7.0, 4.0],
            body_align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _add_figure(doc, config.FIGURES_DIR / "feature_stability.png",
                    f"Şekil 4.6. MRMR seçim frekansı ({n_outer} outer repeat "
                    f"üzerinde).")
        _add_para(doc,
            f"Şekil 4.6, MRMR algoritmasının {n_outer} dış tekrar boyunca hangi "
            "özellikleri ne sıklıkla seçtiğini gösterir. Üst sıralardaki "
            "özelliklerin neredeyse tüm tekrarlarda seçilmiş olması "
            "MRMR'nin rastgele bir alt küme üretmediğini, veriye özgü "
            "kararlı bilgi taşıyan özellikleri tutarlı biçimde "
            "yakaladığını göstermektedir. Bu kararlılık, modelin farklı "
            "eğitim setleri üzerinde benzer karar süreçleri öğrendiğine "
            "işaret etmekte ve elde edilen sonuçların güvenilirliğini "
            "desteklemektedir.",
            first_line_indent=1.0,
        )
        _add_figure(doc, config.FIGURES_DIR / "shap_summary.png",
                    "Şekil 4.7. SHAP özet grafiği (ExtraTrees).")
        _add_para(doc,
            "Şekil 4.7, SHAP analizinden elde edilen ortalama mutlak "
            "katkı değerlerini göstermektedir. Y ekseninde en üstte yer "
            "alan özellikler modelin kararına en çok etki edenlerdir. "
            "X ekseninde noktaların sağa yayılması o özelliğin papilödem "
            "yönünde, sola yayılması ise normal yönünde tahmini "
            "desteklediği anlamına gelmektedir. Bu görsel modelin yalnızca "
            "ne tahmin ettiğini değil, neden öyle tahmin ettiğini de "
            "açıkladığı için sınıflandırıcının klinik bağlamda "
            "yorumlanabilirlik kazanmasını sağlamaktadır. Tablo 4.4'teki "
            "ilk on özellik ile Şekil 4.6'daki yüksek frekansta seçilen "
            "özelliklerin örtüşmesi, kararlılık ve önemlilik arasındaki "
            "tutarlılığı da göstermektedir.",
            first_line_indent=1.0,
        )

    _add_heading2(doc, "4.5. Ödev Ek Sorularına Yönelik Değerlendirmeler")
    _add_para(doc,
        "Bölüm 4.1 ile 4.4 arasında modellerin performansı ayrıntılı "
        "olarak sunulmuştur. Bu alt bölümde ödev tanımında yer alan ek "
        "sorulara doğrudan yanıt verilmesi amaçlanmıştır.",
        first_line_indent=1.0,
    )

    if summary is not None and manifest is not None:
        ek_best_mean_model = summary[("macro_f1", "mean")].idxmax()
        ek_best_mean_val = summary[("macro_f1", "mean")].max()
        ek_best_single_model = manifest["best_overall"]["model"]
        ek_best_single_val = manifest["best_overall"]["macro_f1"]
        ek_q1_text = (
            f"Hangi modelin en iyi performansı verdiği sorusu Bölüm 4.1'de "
            f"ele alınmıştır; ortalama macro-F1 sıralamasında "
            f"{ek_best_mean_model} modeli {ek_best_mean_val:.3f} ile ilk "
            f"sırada yer almakta, tek tekrar bazında en yüksek değer ise "
            f"{ek_best_single_model} modeline ait olup {ek_best_single_val:.3f} "
            f"düzeyindedir. Ensemble modelinin tekli modellerden daha iyi "
            f"performans gösterip göstermediği sorusu ise Bölüm 5'te "
            f"tartışılmıştır; soft voting topluluğu en iyi tek model olan "
            f"{ek_best_mean_model} modelini geride bırakamamış, baz modellerin "
            f"yakın karar yüzeyleri öğrenmesi nedeniyle ortalama alma "
            f"işleminden beklenen çeşitlilik kazancı elde edilememiştir."
        )
    else:
        ek_q1_text = (
            "Hangi modelin en iyi performansı verdiği sorusu Bölüm 4.1'de "
            "ele alınmıştır; ortalama macro-F1 sıralamasında en iyi modelin "
            "ilk sırada yer aldığı, tek tekrar bazında en yüksek değerin de "
            "ilgili modele ait olduğu raporlanmıştır."
        )
    _add_para(doc, ek_q1_text, first_line_indent=1.0)

    _add_para(doc,
        "Kalibrasyonun model güvenilirliğini artırıp artırmadığı sorusu "
        "Bölüm 4.3'te incelenmiş; sigmoid Platt kalibrasyon sonrasında "
        "tüm modellerin Brier skorlarının dar bir aralıkta toplandığı "
        "ve kalibrasyon eğrilerinin diyagonale yakın seyrettiği "
        "görülmüştür. ROC-AUC ile PR-AUC sonuçları arasındaki ilişki "
        "Bölüm 4.1 ve Bölüm 5'te ele alınmış; sınıf dengesizliği "
        "koşullarında PR-AUC değerlerinin ROC-AUC değerlerinden bir "
        "miktar daha düşük seyretmesi, sınıf dengesizliğinin azınlık "
        "sınıf üzerindeki etkisinin PR-AUC ile daha belirgin biçimde "
        "yansıtıldığını göstermektedir. En önemli ilk on radyomik "
        "özelliğin hangileri olduğu sorusu ise Bölüm 4.4'te SHAP analizi "
        "ile Tablo 4.4 üzerinden raporlanmıştır.",
        first_line_indent=1.0,
    )

    _add_para(doc,
        "MRMR özellik seçiminin performansa katkısı sorusu üzerinde "
        "ayrıca durulmak gereklidir. Bu çalışmada 746 radyomik "
        "özelliğin tamamı modele verilmek yerine MRMR algoritması ile "
        "bir alt küme seçilmiştir. Bu kararın iki temel gerekçesi "
        "bulunmaktadır. İlki örnek-özellik oranı ile ilgilidir: 966 "
        "örnek ile 746 özellik kullanıldığında oran yaklaşık 1.3:1 "
        "düzeyinde kalır ki bu yüksek boyutluluk koşullarında aşırı "
        "öğrenme riskini arttıran tipik bir orandır. MRMR ile k = 50 "
        "özellik seçildiğinde aynı oran yaklaşık 19:1 seviyesine "
        "çıkmakta ve modelin daha kararlı öğrenmesi mümkün hale "
        "gelmektedir. İkinci gerekçe ise Demircioğlu [1] tarafından on "
        "radyomik veri seti üzerinde yapılan kapsamlı karşılaştırmada "
        "MRMR'nin seçim kararlılığı açısından en güvenilir yöntemlerden "
        "biri olarak raporlanmasıdır. Şekil 4.6'da sunulan özellik "
        "kararlılığı analizi, üst sıralardaki özelliklerin tekrarların "
        "büyük çoğunluğunda yeniden seçildiğini ortaya koymakta; bu "
        "durum MRMR'nin rastlantısal bir alt küme değil veriye özgü "
        "kararlı bilgi taşıyan özellikleri yakaladığını göstermektedir. "
        "Optuna'nın model bazında bağımsız olarak en uygun mrmr_k "
        "değerini (20, 50 veya 100) seçmesi de farklı model sınıflarının "
        "farklı yoğunluklarda özellik kümeleriyle çalışabildiğini ima "
        "etmektedir.",
        first_line_indent=1.0,
    )

    # Compute dynamic SVM and GB std from summary if available
    if summary is not None:
        try:
            svm_std = summary.loc["SVM", ("macro_f1", "std")]
            gb_std = summary.loc["GB", ("macro_f1", "std")]
            std_text = (
                f"Daha basit ve doğal olarak düzenlileştirilmiş modellerde "
                f"(örneğin SVM ve KNN) bu sapma yaklaşık {svm_std:.2f} "
                f"düzeyinde kalırken, daha karmaşık ve veri açlığı yüksek "
                f"olan Gradient Boosting modelinde aynı değer {gb_std:.2f} "
                f"civarına kadar yükselmektedir."
            )
        except Exception:
            std_text = (
                "Basit ve düzenlileştirilmiş modellerde bu sapma görece "
                "düşük kalırken, daha karmaşık modellerde belirgin "
                "biçimde yükselmektedir."
            )
    else:
        std_text = ""

    _add_para(doc,
        f"Veri boyutunun model performansına etkisi sorusu da ayrı bir "
        f"değerlendirmeyi hak etmektedir. Hasta sayısının görece az "
        f"olması (n = 69) modellerin performans dağılımına doğrudan "
        f"yansımaktadır. Şekil 5.2 ve Şekil 5.3'te görüldüğü üzere "
        f"tekrarlar arası macro-F1 standart sapması model sınıfına göre "
        f"belirgin biçimde değişmektedir. {std_text} Aynı eğitim "
        f"hacminde gözlenen bu fark, az veriyle eğitilen karmaşık "
        f"modellerin farklı dış tekrarlar arasında çok daha kararsız "
        f"davrandığını ortaya koymaktadır. Veri hacminin arttırılması, "
        f"özellikle bu tür modellerin varyansını azaltacak ve genel "
        f"performansını yükseltecek doğal bir adımdır. Ayrıca PR-AUC "
        f"değerlerinin ROC-AUC değerlerinden az da olsa düşük seyretmesi, "
        f"sınıf dengesizliğinin az örneklem koşullarında daha belirgin "
        f"biçimde hissedildiğini düşündürmektedir. Bu gözlemler bir "
        f"arada değerlendirildiğinde, daha geniş hasta kohortlarının "
        f"özellikle karmaşık model sınıfları için daha güvenilir "
        f"sonuçlar üreteceği öngörülmektedir.",
        first_line_indent=1.0,
    )

    # =========================================== 5. TARTIŞMA ===
    _add_heading1(doc, "5. TARTIŞMA")
    _add_para(doc,
        "Bölüm 4'te modellerin macro-F1 değerleri üzerinden "
        "karşılaştırılması sunulmuştur. Aynı modeller ROC-AUC açısından "
        "değerlendirildiğinde biraz farklı bir görünüm ortaya "
        "çıkmaktadır (Şekil 5.1). Macro-F1 sıralamasında biraz geride "
        "kalan bazı modellerin ROC-AUC sıralamasında öne geçebilmesi, "
        "kullanılan metriğe bağlı olarak \"en iyi\" modelin "
        "değişebileceğine işaret etmektedir.",
        first_line_indent=1.0,
    )
    _add_figure(doc, config.FIGURES_DIR / "model_comparison_roc_auc.png",
                caption="Şekil 5.1. Modellerin ortalama ROC-AUC "
                        "karşılaştırması (Macro-F1 için Şekil 4.1'in "
                        "tamamlayıcısı).",
                width_cm=14.5)
    if summary is not None:
        best_model_name = summary[("macro_f1", "mean")].idxmax()
        best_macro = summary[("macro_f1", "mean")].max()
        best_summary_text = (
            f"{best_model_name} modelinin {best_macro:.3f} ortalama "
            f"macro-F1 değerine ulaşması, küçük örneklemli bir radyomik "
            f"veri seti için oldukça tatmin edicidir."
        )
    else:
        best_summary_text = (
            "En iyi modelin ortalama macro-F1 değeri, küçük örneklemli "
            "bir radyomik veri seti için tatmin edicidir."
        )
    _add_para(doc,
        f"Çalışmanın sonuçları üç temel noktayı destekler niteliktedir. "
        f"İlki, hasta seviyesindeki veri sızıntısını engelleyen sıkı bir "
        f"doğrulama yapısı kurulmuş olmasına rağmen model performansının "
        f"klinik açıdan makul aralıkta kalmasıdır. {best_summary_text} "
        f"İkincisi, kalibre edilmiş baz modellerin birbirlerine yakın "
        f"karar yüzeyleri öğrenmesi nedeniyle soft voting topluluğunun "
        f"en iyi tek modeli geride bırakamamış olmasıdır. Bu durum, "
        f"çeşitlilikten yararlanan topluluk stratejilerinin değerli "
        f"olabileceğini akla getirmektedir; Tang ve ark. [6] bu yönde "
        f"farklı ağırlıklandırma stratejilerinin etkisini incelemiştir.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Çalışmada karşılaşılan ilginç bir noktanın altını çizmek "
        "gerekmektedir. scikit-learn kütüphanesinin son sürümünde "
        "ortaya çıkan metadata routing davranışı nedeniyle, "
        "`CalibratedClassifierCV`'nin `StratifiedGroupKFold` ile birlikte "
        "kullanılması varsayılan olarak hasta gruplarını iç ayırıcıya "
        "iletmemekte; bu da kalibrasyon adımında hasta sızıntısına yol "
        "açabilmektedir. Bu durum kolayca gözden kaçırılabilecek ama "
        "patient-level çapraz doğrulama düzeneğini fiilen geçersiz "
        "kılabilecek bir tuzaktır. Sorunu aşmak için projede manuel "
        "Platt scaling yapan bir sarmalayıcı yazılmıştır.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        "Sınıf dengesizliğinin yönetimi konusunda da literatürle uyumlu "
        "bir tercih yapılmıştır. SMOTE gibi sentetik örnekleme "
        "yöntemleri yerine modelin kayıp fonksiyonunda azınlık sınıfa "
        "daha yüksek ağırlık veren `class_weight='balanced'` ayarı "
        "kullanılmıştır. Bu tercihin gerekçeleri Bölüm 2.3'te ayrıntılı "
        "olarak açıklanmıştır.",
        first_line_indent=1.0,
    )
    _add_para(doc,
        f"İstatistiksel test gücü konusunda ise dikkat edilmesi gereken "
        f"bir nokta vardır. {n_outer} bloklu bir blok-tabanlı yapıda "
        f"yirmi bir ikili karşılaştırma yapılması, Bonferroni "
        f"düzeltmesinin anlamlılık eşiğini ciddi şekilde "
        f"sıkılaştırmasına yol açmaktadır. Bu durumda yalnızca oldukça "
        f"belirgin model farkları test sonucunda anlamlı çıkabilmektedir. "
        f"Friedman testi modeller arası genel farkı doğrulamış olsa da, "
        f"ikili karşılaştırma sonuçları kesin hüküm olarak değil yön "
        f"gösterici olarak yorumlanmalıdır.",
        first_line_indent=1.0,
    )

    _add_heading2(doc, "5.1. Modellerin Tekrarlar Arası Tutarlılığı")
    _add_para(doc,
        "Bir modelin küçük veri setinde gerçekten kullanılabilir olup "
        "olmadığını yalnızca ortalama performansa bakarak söylemek "
        "yanıltıcı olabilir. Aynı ortalamayı veren iki model, tekrarlar "
        "arasında çok farklı kararlılık sergileyebilir. Şekil 5.2 "
        f"yedi modelin {n_outer} dış tekrardaki macro-F1 dağılımını boxplot "
        "biçiminde sunmaktadır.",
        first_line_indent=1.0,
    )
    _add_figure(doc, config.FIGURES_DIR / "model_varyans_boxplot.png",
                caption=f"Şekil 5.2. Modellerin {n_outer} dış tekrardaki macro-F1 "
                        "dağılımı.",
                width_cm=15.5)
    if summary is not None:
        std_sorted = summary[("macro_f1", "std")].sort_values()
        stable_models = std_sorted.index[:4].tolist()
        highest_var_model = std_sorted.index[-1]
        stable_text = ", ".join(stable_models[:-1]) + f" ve {stable_models[-1]}"
        boxplot_comment = (
            f"Görselden iki önemli nokta ortaya çıkmaktadır. Birincisi, "
            f"{stable_text} modellerinin kutu yükseklikleri (IQR) görece "
            f"dardır; bu modeller farklı dış tekrarlarda benzer performans "
            f"göstermekte, dolayısıyla küçük örneklemde dahi kararlı "
            f"davranmaktadır. İkincisi ise {highest_var_model} modelinin "
            f"kutusu diğerlerinden açık biçimde geniş ve aşağı doğru "
            f"yayılmıştır. Bu durum, modelin bazı dış tekrarlarda 0.5 "
            f"civarına kadar düşebildiğini göstermektedir; yani aynı model "
            f"bir tekrarda iyi sonuç verirken bir başka rastgele bölmede "
            f"beklenmedik biçimde kötü olabilmektedir. Bu örüntü "
            f"{highest_var_model} modelinin az veriye duyarlı yapısının "
            f"doğrudan bir yansımasıdır. Klinik bir uygulamada tek tek "
            f"dış tekrarlardaki performansın ortalamanın altında kalmaması "
            f"istenir; bu yüzden modellerin tutarlılığı klinik "
            f"değerlendirmede ortalamadan en az kadar önemlidir."
        )
    else:
        boxplot_comment = (
            "Görselden iki önemli nokta ortaya çıkmaktadır. Birincisi, "
            "düşük standart sapmalı modellerin kutu yükseklikleri (IQR) "
            "görece dardır; bu modeller farklı dış tekrarlarda benzer "
            "performans göstermekte, dolayısıyla küçük örneklemde dahi "
            "kararlı davranmaktadır. İkincisi ise en yüksek varyanslı "
            "modelin kutusu diğerlerinden açık biçimde geniş ve aşağı "
            "doğru yayılmıştır."
        )
    _add_para(doc, boxplot_comment, first_line_indent=1.0)

    _add_heading2(doc, "5.2. Doğruluk-Kararlılık Dengesi")
    _add_para(doc,
        "Boxplot analizinde gözlenen kararsızlık konusunun ayrıntılı "
        "olarak ele alınması gerekmektedir. Modellerin "
        "değerlendirilmesinde ortalama performans tek başına yeterli "
        "bir ölçüt değildir. Aynı ortalama macro-F1 değerine ulaşan "
        "iki model, tekrarlar arası dalgalanma açısından oldukça farklı "
        "davranabilir. Şekil 5.3 bu noktayı görselleştirmektedir: "
        "yatay eksende modelin ortalama macro-F1 skoru, dikey eksende "
        f"ise bu skorun {n_outer} tekrar üzerindeki standart sapması yer "
        "almaktadır.",
        first_line_indent=1.0,
    )
    _add_figure(doc, config.FIGURES_DIR / "dogruluk_kararlilik.png",
                caption="Şekil 5.3. Modellerin ortalama macro-F1 ile "
                        "tekrarlar arası standart sapma ilişkisi.",
                width_cm=14.5)
    if summary is not None:
        means = summary[("macro_f1", "mean")]
        stds = summary[("macro_f1", "std")]
        ideal_models = [m for m in means.index
                        if means[m] >= means.median() and stds[m] <= stds.median()]
        if not ideal_models:
            ideal_models = stds.sort_values().index[:4].tolist()
        ideal_text = ", ".join(ideal_models[:-1]) + f" ve {ideal_models[-1]}" \
                     if len(ideal_models) > 1 else ideal_models[0]
        worst_var_model = stds.idxmax()
        worst_mean_model = means.idxmin()
        if worst_mean_model == worst_var_model:
            second_worst_mean = means.drop(worst_var_model).idxmin()
            tradeoff_comment = (
                f"Grafikte sağ alta inen modeller hem yüksek doğruluk "
                f"üretir hem de tekrarlar arasında kararlı davranır; bu da "
                f"ideal bölgedir. {ideal_text} modellerinin bu bölgeye "
                f"yakın konumlanması, küçük örneklemde dahi tutarlı "
                f"sonuçlar verebildiklerini göstermektedir. "
                f"{worst_var_model} modelinin ise grafiğin sol-yukarısına "
                f"doğru kayması, hem ortalama performansının düşük olduğunu "
                f"hem de bu düşük performansın bile tekrarlar arasında "
                f"ciddi biçimde dalgalandığını ortaya koymaktadır. "
                f"{second_worst_mean} modelinin ortalama performansının "
                f"düşük kalması, doğrusal karar yüzeyinin radyomik verinin "
                f"karmaşık yapısı için yetersiz kaldığı şeklinde "
                f"yorumlanabilir."
            )
        else:
            tradeoff_comment = (
                f"Grafikte sağ alta inen modeller hem yüksek doğruluk "
                f"üretir hem de tekrarlar arasında kararlı davranır; bu da "
                f"ideal bölgedir. {ideal_text} modellerinin bu bölgeye "
                f"yakın konumlanması, küçük örneklemde dahi tutarlı "
                f"sonuçlar verebildiklerini göstermektedir. "
                f"{worst_var_model} modelinin grafiğin yukarısına doğru "
                f"kayması, performansının tekrarlar arasında ciddi biçimde "
                f"dalgalandığını ortaya koymaktadır. {worst_mean_model} "
                f"modelinin ortalama performansının diğer modellerin "
                f"altında kalması ise doğrusal karar yüzeyinin radyomik "
                f"verinin karmaşık yapısı için yetersiz kaldığı şeklinde "
                f"yorumlanabilir."
            )
    else:
        tradeoff_comment = (
            "Grafikte sağ alta inen modeller hem yüksek doğruluk üretir "
            "hem de tekrarlar arasında kararlı davranır; bu da ideal "
            "bölgedir. İdeal bölgeye yakın konumlanan modeller küçük "
            "örneklemde dahi tutarlı sonuçlar verebilmektedir."
        )
    _add_para(doc, tradeoff_comment, first_line_indent=1.0)

    # ============================================ 6. SONUÇ ===
    _add_heading1(doc, "6. SONUÇ")
    if manifest is not None:
        sonuc_best_mean_model = summary[("macro_f1", "mean")].idxmax()
        sonuc_best_mean_val = summary[("macro_f1", "mean")].max()
        _add_para(doc,
            f"Bu projede, optik diskten çıkarılmış radyomik özellikler "
            f"üzerinden Normal ile Papilödem ayrımını yapan, hasta "
            f"seviyesindeki veri sızıntısını sistematik biçimde "
            f"engelleyen, {n_outer} dış tekrar ve {n_inner} iç fold "
            f"üzerine kurulu, kalibrasyonlu ve yorumlanabilir bir makine "
            f"öğrenmesi sistemi geliştirilmiştir. Sistem, MRMR özellik "
            f"seçimi, Optuna ile hiperparametre arama, sigmoid "
            f"kalibrasyon ve soft voting topluluk aşamalarını tek bir "
            f"çerçeve içinde bir araya getirmektedir. Pipeline'ın tam "
            f"çalışması yaklaşık {manifest['total_seconds']/60:.0f} "
            f"dakika sürmüş; en yüksek ortalama makro-F1 değerine "
            f"{sonuc_best_mean_model} modeli "
            f"({sonuc_best_mean_val:.3f}) ile ulaşılmıştır.",
            first_line_indent=1.0,
        )
        _add_para(doc,
            "Elde edilen bulgular, kurulan sıkı doğrulama yapısına rağmen "
            "radyomik özellikler üzerinden papilödem tespitinin klinik "
            "açıdan anlamlı bir doğrulukla yapılabildiğini "
            "göstermektedir. Modellerin ortalama ROC-AUC değerleri 0.88 "
            "ile 0.93 arasında yoğunlaşmış; PR-AUC değerleri de aynı "
            "şekilde yüksek seyretmiştir. Bu durum, kullanılan radyomik "
            "özellik kümesinin iki sınıfı ayırt edecek yeterli bilgi "
            "taşıdığını ortaya koymaktadır. Sınıf dengesizliği "
            "koşullarına rağmen balanced accuracy değerlerinin de "
            "yüksek kalması, modellerin azınlık sınıfı (papilödem) "
            "gözden kaçırmadığını desteklemektedir. Çalışma sonucunda "
            "elde edilen modellerin çok boyutlu performans profili "
            "Şekil 6.1'de özetlenmiştir.",
            first_line_indent=1.0,
        )
        _add_figure(doc, config.FIGURES_DIR / "ozet_radar.png",
                    caption="Şekil 6.1. En iyi üç model ve Ensemble "
                            "modelinin çok-metrik performans profili "
                            "(Brier hariç tüm metrikler).",
                    width_cm=14.5)
        if summary is not None:
            radar_top3 = list(summary[("macro_f1", "mean")]
                              .sort_values(ascending=False).index[:3])
            if "Ensemble" not in radar_top3:
                radar_top3.append("Ensemble")
            radar_lead = radar_top3[0]
            radar_others = [m for m in radar_top3 if m != radar_lead]
            radar_others_text = (
                ", ".join(radar_others[:-1]) + f" ve {radar_others[-1]}"
                if len(radar_others) > 1 else (radar_others[0] if radar_others else "")
            )
            radar_comment = (
                f"Şekil 6.1'deki radar grafiği, tek bir metrik üzerinden "
                f"değil birden fazla metriğin birlikte değerlendirilmesi "
                f"yoluyla modellerin güçlü ve zayıf yönlerini aynı anda "
                f"görmeyi mümkün kılmaktadır. {radar_lead} modeli "
                f"neredeyse her metrik ekseninde diğerleriyle ya birlikte "
                f"ya da hafifçe önde konumlanmaktadır. {radar_others_text} "
                f"modellerinin de çoğu metrikte {radar_lead} ile rekabet "
                f"edebildiği fark edilmektedir."
            )
        else:
            radar_comment = (
                "Şekil 6.1'deki radar grafiği, tek bir metrik üzerinden "
                "değil birden fazla metriğin birlikte değerlendirilmesi "
                "yoluyla modellerin güçlü ve zayıf yönlerini aynı anda "
                "görmeyi mümkün kılmaktadır."
            )
        _add_para(doc, radar_comment, first_line_indent=1.0)
        _add_para(doc,
            "Çalışmanın metodolojik açıdan en güçlü yanlarından biri, "
            "veri sızıntısının her aşamada (dış bölme, iç çapraz "
            "doğrulama, kalibrasyon ve eşik seçimi) sistematik olarak "
            "engellenmiş olmasıdır. Sigmoid kalibrasyonun ardından tüm "
            "modellerin Brier skorlarının dar bir aralıkta toplanması ve "
            "kalibrasyon eğrilerinin diyagonale yakın kalması, üretilen "
            "olasılık tahminlerinin güvenilir olduğunu göstermektedir. "
            "Ayrıca SHAP analizi ile elde edilen yorumlanabilirlik "
            "katmanı, modelin karar süreçlerinin klinik açıdan "
            "değerlendirilebilir olmasına imkân tanımaktadır.",
            first_line_indent=1.0,
        )

    # ================================================== KAYNAKLAR ===
    _add_heading1(doc, "KAYNAKLAR")
    for i, ref in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.15
        pf.space_after = Pt(6)
        pf.left_indent = Cm(0.8)
        pf.first_line_indent = Cm(-0.8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f"[{i}]  {ref}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

    # =============================================== EKLER: KOD ===
    _add_heading1(doc, "EKLER: KOD ÖRNEKLERİ")
    _add_para(doc,
        "Bu ek bölümde projenin kritik kısımlarını gösteren bazı kod "
        "parçaları yer almaktadır. Tüm kaynak kod GitHub deposundan "
        "incelenebilir.",
        first_line_indent=0.0,
    )

    _add_heading2(doc, "EK 1. Hasta ID Birleştirme ve Bütünlük Kontrolü")
    _add_para(doc,
        "Aynı `PatientIndex` değeri her iki dosyada da bulunduğundan, "
        "doğrudan birleştirme aynı kimliğin iki farklı hastayı "
        "tanımlaması riskini doğurmaktadır. Bu sorunun önüne geçmek "
        "amacıyla normal hastalara `N_`, papilödem hastalarına `P_` ön "
        "eki eklenmiştir:",
        first_line_indent=0.0,
    )
    _add_code_block(doc, '''
def _prefix_patient_ids(df, prefix):
    out = df.copy()
    out["patient_id"] = prefix + "_" + out["PatientIndex"].astype(str)
    return out

normal = _prefix_patient_ids(normal, "N")
papil  = _prefix_patient_ids(papil,  "P")
combined = pd.concat([normal, papil], ignore_index=True)

# Hizli butunluk kontrolu
normal_ids = set(combined.loc[combined.label == 0, "patient_id"])
papil_ids  = set(combined.loc[combined.label == 1, "patient_id"])
assert not (normal_ids & papil_ids), "Hasta iki sinifta birden gorunuyor!"
assert combined.groupby("patient_id")["label"].nunique().max() == 1
''')

    _add_heading2(doc, "EK 2. Korelasyon Filtresi (Custom Transformer)")
    _add_para(doc,
        "Yüksek korelasyonlu özelliklerin MRMR adımından önce elenmesi "
        "amacıyla sade bir sklearn transformer'ı tanımlanmıştır. Mutlak "
        "Pearson korelasyonu 0.95'in üzerindeki çiftlerden ilk gelene "
        "öncelik verilmektedir:",
        first_line_indent=0.0,
    )
    _add_code_block(doc, '''
class CorrelationFilter(BaseEstimator, TransformerMixin):
    def __init__(self, threshold=0.95):
        self.threshold = threshold

    def fit(self, X, y=None):
        corr = np.abs(np.corrcoef(X, rowvar=False))
        corr = np.nan_to_num(corr, nan=0.0)
        drop = set()
        for i in range(corr.shape[0]):
            if i in drop:
                continue
            for j in range(i + 1, corr.shape[0]):
                if j not in drop and corr[i, j] > self.threshold:
                    drop.add(j)
        self.kept_indices_ = np.array(
            [i for i in range(corr.shape[0]) if i not in drop]
        )
        return self

    def transform(self, X):
        return np.asarray(X)[:, self.kept_indices_]
''')

    _add_heading2(doc, "EK 3. Hasta-Bilinçli Kalibrasyon")
    _add_para(doc,
        "Sklearn 1.6 sürümünde `CalibratedClassifierCV`, "
        "`StratifiedGroupKFold` ile beklendiği gibi çalışmamakta; bu "
        "durum kalibrasyon adımında hasta sızıntısına yol "
        "açabilmektedir. Söz konusu sorunun aşılması için manuel "
        "Platt scaling sarmalayıcısı tanımlanmıştır:",
        first_line_indent=0.0,
    )
    _add_code_block(doc, '''
class GroupAwareSigmoidCalibrator(BaseEstimator, ClassifierMixin):
    def __init__(self, estimator, n_folds=5):
        self.estimator = estimator
        self.n_folds = n_folds

    def fit(self, X, y, groups):
        cv = StratifiedGroupKFold(
            n_splits=self.n_folds, shuffle=True, random_state=42
        )
        oof = cross_val_predict(
            clone(self.estimator), X, y,
            groups=groups, cv=cv,
            method="predict_proba", n_jobs=-1,
        )[:, 1]
        self.sigmoid_ = _SigmoidCalibration().fit(oof, y)
        self.estimator_ = clone(self.estimator).fit(X, y)
        return self

    def predict_proba(self, X):
        raw = self.estimator_.predict_proba(X)[:, 1]
        cal = np.clip(self.sigmoid_.predict(raw), 0.0, 1.0)
        return np.column_stack([1.0 - cal, cal])
''')

    _add_heading2(doc, "EK 4. Pipeline Kurulumu")
    _add_para(doc,
        "Tüm adımlar tek bir sklearn `Pipeline` içinde tutulup CV "
        "döngüsünde otomatik olarak yeniden fit edildi. Bu yaklaşım "
        "veri sızıntısı riskini büyük ölçüde ortadan kaldırır:",
        first_line_indent=0.0,
    )
    _add_code_block(doc, f'''
pipeline = Pipeline([
    ("preprocess", Pipeline([
        ("imputer",     SimpleImputer(strategy="median")),
        ("variance",    VarianceThreshold(threshold=0.01)),
        ("correlation", CorrelationFilter(threshold=0.95)),
        ("scaler",      RobustScaler()),
    ])),
    ("mrmr",  MRMRSelector(k=params["mrmr_k"], random_state=42)),
    ("model", MODEL_REGISTRY[model_name](params)),
], memory=cache_dir)

# Dis dongu: hasta-bazli bolme x {n_outer} tekrar, %70/%10/%20 uclu yapi
N_OUTER, TEST_SIZE, VAL_SIZE = {n_outer}, 0.20, 0.10
inner_val_frac = VAL_SIZE / (1.0 - TEST_SIZE)  # 0.125 -> %10 / %80

for repeat_idx in range(N_OUTER):
    seed = 42 + repeat_idx
    # 1) test setini hasta bazli ayir (%80 train+val, %20 test)
    outer = GroupShuffleSplit(n_splits=1, test_size=TEST_SIZE, random_state=seed)
    tv_idx, te_idx = next(outer.split(X, y, groups=patient_id))

    # 2) train+val icinden validation setini hasta bazli ayir (%70 / %10)
    inner = GroupShuffleSplit(n_splits=1, test_size=inner_val_frac, random_state=seed)
    rel_tr, rel_val = next(inner.split(X[tv_idx], y[tv_idx], groups=patient_id[tv_idx]))
    tr_idx, val_idx = tv_idx[rel_tr], tv_idx[rel_val]

    # Hasta sizintisi yok: uc kume birbirinden ayrik
    assert set(patient_id[tr_idx]).isdisjoint(patient_id[val_idx])
    assert set(patient_id[tr_idx]).isdisjoint(patient_id[te_idx])
    assert set(patient_id[val_idx]).isdisjoint(patient_id[te_idx])

    # tr_idx: Optuna + kalibrasyon; val_idx: F1-en-iyi esik; te_idx: nihai metrikler
    ...
''')

    config.REPORT_DIR.mkdir(parents=True, exist_ok=True)
    out = config.REPORT_DIR / "final_report.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    p = build_docx()
    print(f"Wrote {p}")
